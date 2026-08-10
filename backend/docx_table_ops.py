"""Raw OOXML column insertion for python-docx tables -- python-docx itself has no
public API for adding a table column, so this manipulates the underlying lxml
elements directly: the table's <w:tblGrid> (column-width grid) and each row's
cell elements. Isolated in its own small module since this is genuine
document-structure surgery (unlike the rest of this codebase, which only ever
writes text into pre-existing cells) -- used by clips_nonpkpd_parser.py /
populate_rfp.py to give two files assigned to the same Specimen Management /
Referral Lab column ("LTS Serum", say) each their own real "LTS Serum (1)" /
"LTS Serum (2)" column instead of overwriting one another.

Also exports row_cells(), a corrected replacement for python-docx's own
`Row.cells` -- confirmed via direct XML inspection that several data rows in the
template's Specimen Management / Referral Lab tables wrap some of their cells in
a Word content control (`<w:sdt><w:sdtContent><w:tc>...`) for a dropdown-style
cell, and python-docx's `Row.cells` only finds *plain* `<w:tc>` children of the
row, silently skipping any sdt-wrapped ones. A row can mix both (e.g. one
row: 3 plain <w:tc> then 3 sdt-wrapped ones) -- using `Row.cells` on such a row
undercounts it and throws off every position-based column lookup, which is
exactly what caused the "Analyte name" row's LTS PK cell to appear unreachable
before this fix (it wasn't actually merged at all).
"""
from __future__ import annotations

import copy

from lxml import etree

from docx.table import _Cell

_W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
_W_TC = f'{_W_NS}tc'
_W_SDT = f'{_W_NS}sdt'
_W_SDT_CONTENT = f'{_W_NS}sdtContent'
_W_R_PR = f'{_W_NS}rPr'


def _row_containers(row):
    """Returns [(container_element, tc_element), ...] for every real per-grid-
    column cell in `row`, in document order. `container_element` is what must be
    cloned/moved to duplicate or reorder the cell as a whole -- the <w:tc> itself
    for a plain cell, or the enclosing <w:sdt> for a content-control-wrapped one
    (cloning just the inner <w:tc> would leave two <w:tc>s inside one
    <w:sdtContent>, which isn't valid). `tc_element` is always the actual
    <w:tc>, used for merge/text inspection regardless of wrapping."""
    out = []
    for child in row._tr:
        if child.tag == _W_TC:
            out.append((child, child))
        elif child.tag == _W_SDT:
            content = child.find(_W_SDT_CONTENT)
            tc = content.find(_W_TC) if content is not None else None
            if tc is not None:
                out.append((child, tc))
    return out


def row_cells(row):
    """Like python-docx's `Row.cells`, but also finds cells wrapped in a content
    control (see module docstring) -- returns real docx.table._Cell objects for
    every per-grid-column cell in `row`, in document order, so `.text`/
    `.paragraphs`/etc. all work normally on the result."""
    return [_Cell(tc, row.table) for _container, tc in _row_containers(row)]


def _has_merge(tc) -> bool:
    tc_pr = tc.find(f'{_W_NS}tcPr')
    if tc_pr is None:
        return False
    grid_span = tc_pr.find(f'{_W_NS}gridSpan')
    if grid_span is not None and int(grid_span.get(f'{_W_NS}val', '1')) > 1:
        return True
    return tc_pr.find(f'{_W_NS}vMerge') is not None


def _set_header_text(cell, text: str) -> None:
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first_para = cell.paragraphs[0]
    for run in list(first_para.runs):
        run._element.getparent().remove(run._element)
    first_para.add_run(text)


def insert_column_after(table, col_index: int, header_text: str) -> int:
    """Duplicates the column at `col_index` (0-indexed), inserting the clone
    immediately to its right, for every row in `table` (header row included) --
    and the matching <w:gridCol> in the table's <w:tblGrid>. Whichever container
    a row's target cell actually uses (a plain <w:tc>, or a content-control
    <w:sdt> wrapping one -- see row_cells()'s docstring) is cloned as a whole, so
    a duplicated dropdown cell stays a real, working dropdown. The new header
    cell's text is replaced with `header_text`; every other (data) row's cloned
    cell starts as a verbatim copy of the original column's cell at that row
    (same formatting, same text if any -- normally blank/placeholder for a
    freshly-templated column), ready for a caller to overwrite with real values
    afterward.

    Returns the new column's 0-indexed position (always col_index + 1).

    Rows that are already fully merged across the whole table width at this
    point (e.g. a free-text "shipping instructions" row spanning every column as
    one cell -- confirmed present in the real template's Referral Lab / Storage
    Samples tables, for rows that never hold per-column data anyway) have no
    distinct cell at `col_index` to duplicate at all; those rows are left
    untouched rather than raising, since inserting a column doesn't need to
    touch them -- fill_spec_by_index()'s own existing merged-cell fallback
    (appending a bracketed note into cell 0) already handles writing into such a
    row for any column, same as before this function existed.

    Raises ValueError if a row genuinely has a *distinct* cell at col_index
    (i.e. its own per-column cell, not lumped into an all-column merge) that is
    itself part of a horizontal merge (gridSpan > 1) or vertical merge (vMerge)
    -- inserting a column next to a partial merge like that would corrupt the
    table's grid alignment, and no column this is ever called on in this
    codebase's real usage has that shape (confirmed via direct inspection of the
    template's Referral Lab / Storage Samples data rows) -- a raised,
    clearly-worded error here beats silently producing a malformed .docx if that
    assumption is ever violated by a template change.
    """
    tbl = table._tbl

    for row in table.rows:
        containers = _row_containers(row)
        if col_index >= len(containers):
            continue  # fully merged row at this position -- nothing to duplicate
        _container, tc = containers[col_index]
        if _has_merge(tc):
            raise ValueError(f'cannot insert a column next to a merged cell at col_index {col_index}')

    grid = tbl.find(f'{_W_NS}tblGrid')
    if grid is not None:
        grid_cols = grid.findall(f'{_W_NS}gridCol')
        if col_index < len(grid_cols):
            grid_cols[col_index].addnext(copy.deepcopy(grid_cols[col_index]))

    for row in table.rows:
        containers = _row_containers(row)
        if col_index >= len(containers):
            continue  # fully merged row at this position -- left untouched, see docstring
        container, _tc = containers[col_index]
        container.addnext(copy.deepcopy(container))

    new_index = col_index + 1
    _set_header_text(row_cells(table.rows[0])[new_index], header_text)
    return new_index


def delete_columns(table, col_indices) -> None:
    """Removes the columns at `col_indices` (0-indexed) entirely -- the matching
    <w:gridCol> from <w:tblGrid>, and each row's container (a plain <w:tc>, or
    the wrapping <w:sdt> for a content-control cell -- see row_cells()'s own
    docstring) at that position. Inverse of insert_column_after().

    Rows already fully merged at a given index (fewer containers than the
    column count -- e.g. a free-text "shipping instructions" row spanning
    every column as one cell) simply have nothing to remove there and are left
    untouched, same convention insert_column_after() already uses.

    `col_indices` may be given in any order -- they're processed internally in
    descending order so removing one column never shifts the position of
    another one still queued for removal. Duplicate indices are ignored."""
    tbl = table._tbl
    grid = tbl.find(f'{_W_NS}tblGrid')
    grid_cols = grid.findall(f'{_W_NS}gridCol') if grid is not None else []

    for col_index in sorted(set(col_indices), reverse=True):
        if col_index < len(grid_cols):
            grid_col = grid_cols[col_index]
            grid_col.getparent().remove(grid_col)
            grid_cols.pop(col_index)

        for row in table.rows:
            containers = _row_containers(row)
            if col_index >= len(containers):
                continue  # fully merged row at this position -- nothing to remove
            container, _tc = containers[col_index]
            container.getparent().remove(container)


def _xml_escape(text: str) -> str:
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def strip_dropdown_and_set_text(cell, text: str) -> None:
    """Removes an inline dropdown content control (a <w:sdt> with a
    <w:dropDownList>, placed inside the cell's own paragraph -- NOT the
    row-level cell-wrapping <w:sdt> row_cells()/insert_column_after() handle)
    from `cell` and replaces it with a plain run containing `text`.

    Unlike populate_rfp._write_cell_value()'s append-beside-the-dropdown
    behavior (used for the Referral Lab's Analyte-name row historically), this
    actually deletes the picklist -- per direct confirmation that the rows
    this is used for (Analyte name's LTS PK/Immunogenicity cells, Validated
    assay, Ref Lab contract owner, Special collection tube required?) should
    show ONLY the extracted/regex/hardcoded value, not a dropdown plus text
    beside it.

    Falls back to a plain overwrite of the cell's first paragraph if no
    dropdown is present at all (e.g. the bmkr/Storage Samples columns'
    Analyte name cells, which are already plain text with no picklist)."""
    sdt = cell._tc.find('.//' + _W_SDT)
    if sdt is None:
        _plain_overwrite(cell, text)
        return

    # Reuse the dropdown's own run formatting (font/size/color) for visual
    # continuity, when its placeholder content has any.
    content = sdt.find(_W_SDT_CONTENT)
    rpr_el = content.find('.//' + _W_R_PR) if content is not None else None
    rpr_xml = etree.tostring(rpr_el, encoding='unicode') if rpr_el is not None else ''

    new_run = etree.fromstring(
        f'<w:r xmlns:w="{_W_NS[1:-1]}">{rpr_xml}'
        f'<w:t xml:space="preserve">{_xml_escape(text)}</w:t></w:r>'
    )
    sdt.addprevious(new_run)
    sdt.getparent().remove(sdt)


def _plain_overwrite(cell, text: str) -> None:
    """Minimal plain-text overwrite of a cell's first paragraph, self-contained
    here (rather than importing populate_rfp.set_cell_text) to avoid a
    circular import between this module and populate_rfp.py."""
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ''
