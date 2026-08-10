#!/usr/bin/env python3
"""Central Laboratory RFP auto-populator (OIAH).

Reads the OIAH protocol + clinical design-element markdowns, derives RFP field
values, and writes them into the blank .docx template IN PLACE (preserving the
table layout and formatting). Produces:
  - output/Central_Laboratory_RFP_OIAH_populated.docx
  - output/RFP_fill_report.md   (every field: value, source, status)

Rules:
  * No value is invented. Anything not located in a source is written as a
    visible review token and logged as 'review' in the report.
  * Source precedence: protocol = clinical identity / analytes; design = study
    ops / design flags; profile = requestor; computed = derived dates.
"""
import logging, re, datetime, os, json, sys
from dataclasses import dataclass, field
from typing import Optional, List
from pathlib import Path
import docx
from docx.oxml.ns import qn
from docx.shared import RGBColor

logger = logging.getLogger(__name__)

from extractors import (protocol_number, compound, protocol_title, phase,
    therapeutic_area, indication, min_age, is_oncology,
    immunogenicity, genetics_pgx, screen_fail_rate, extract_country_check_row,
    COUNTRY_LANG, REVIEW, Finding, parse_manual_countries)
from design_parser import DesignParser
from md_table import (find_heading_by_number, extract_section,
                      parse_soa_text, parse_analytes_text,
                      find_pipe_tables, parse_pipe_table,
                      is_non_lab, is_category_header, is_visit_column)


@dataclass
class Report:
    """Result of a populate_rfp run."""
    findings: list[Finding] = field(default_factory=list)
    filled: int = 0
    computed: int = 0
    review_count: int = 0
    report_text: str = ''
    output_path: str = ''
    # Structured data for API/UI consumers that want it without re-parsing
    # formatted finding strings (e.g. a web UI's field/table display).
    countries: list = field(default_factory=list)
    soa_headers: list = field(default_factory=list)
    soa_rows: list = field(default_factory=list)


def strip_json_suffix(text: str) -> str:
    """Strip JSON appendix (pages array) appended by liteparse/PyMuPDF output.

    The JSON block is always at the END of the text and starts with a line
    that is exactly '{' followed by '"pages"'. We search backwards from the end
    to avoid false matches on '{' characters in the document body.
    """
    lines = text.splitlines()
    # Search backwards for the JSON boundary — it's always near the end
    for i in range(len(lines) - 1, max(5, len(lines) - 5000), -1):
        if lines[i].strip() == '{':
            # Verify it's actually the start of our JSON pages block
            remaining = '\n'.join(lines[i:i+3])
            if '"pages"' in remaining:
                return '\n'.join(lines[:i])
    return text


# ---------------------------------------------------------------- docx helpers
def distinct(row):
    out, seen = [], set()
    for c in row.cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc)); out.append(c)
    return out

def append_val(cell, text):
    p = cell.paragraphs[0]
    base = p.runs[0] if p.runs else None
    run = p.add_run(('  ' if (base and base.text and not base.text.endswith(' ')) else '') + text)
    if base is not None:
        run.bold = base.bold
        if base.font.size: run.font.size = base.font.size
        if base.font.name: run.font.name = base.font.name
    return run

def set_content_control(cell, value):
    """Replace a Word content-control (w:sdt) placeholder in `cell` with `value`,
    recoloured to black and with the placeholder flag cleared."""
    for sdt in cell._tc.iter(qn('w:sdt')):
        content = sdt.find(qn('w:sdtContent'))
        if content is None:
            continue
        ts = content.findall('.//' + qn('w:t'))
        if not ts:
            continue
        ts[0].text = value
        ts[0].set(qn('xml:space'), 'preserve')
        for t in ts[1:]:
            t.text = ''
        for col in content.iter(qn('w:color')):
            col.set(qn('w:val'), '000000')
        pr = sdt.find(qn('w:sdtPr'))
        if pr is not None:
            for tag in ('w:placeholder', 'w:showingPlcHdr'):
                el = pr.find(qn(tag))
                if el is not None:
                    pr.remove(el)
        return True
    return False

def set_nth_content_control(cell, n, value):
    """Like set_content_control, but targets the Nth (0-indexed) w:sdt in `cell` specifically --
    needed when a single cell holds more than one content control (e.g. two yes/no questions
    packed into one cell) and each must be filled with a different, independent answer."""
    sdts = list(cell._tc.iter(qn('w:sdt')))
    if n >= len(sdts):
        return False
    sdt = sdts[n]
    content = sdt.find(qn('w:sdtContent'))
    if content is None:
        return False
    ts = content.findall('.//' + qn('w:t'))
    if not ts:
        return False
    ts[0].text = value
    ts[0].set(qn('xml:space'), 'preserve')
    for t in ts[1:]:
        t.text = ''
    for col in content.iter(qn('w:color')):
        col.set(qn('w:val'), '000000')
    pr = sdt.find(qn('w:sdtPr'))
    if pr is not None:
        for tag in ('w:placeholder', 'w:showingPlcHdr'):
            el = pr.find(qn(tag))
            if el is not None:
                pr.remove(el)
    return True

def set_cell_text(cell, text):
    if set_content_control(cell, text):
        return
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        try: p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        except Exception as _e: logger.warning('set_cell_text color: %s', _e)
        for r in p.runs[1:]: r.text = ''
    else:
        r = p.add_run(text)
        try: r.font.color.rgb = RGBColor(0, 0, 0)
        except Exception as _e: logger.warning('set_cell_text add_run color: %s', _e)
    for ex in cell.paragraphs[1:]:
        for r in ex.runs: r.text = ''


def _append(cell, text):
    p = cell.paragraphs[0]
    base = p.runs[0] if p.runs else None
    run = p.add_run('  ' + text)
    if base is not None:
        run.bold = base.bold
        if base.font.size: run.font.size = base.font.size
        if base.font.name: run.font.name = base.font.name


def fill(tbl_idx, row, label_sub, value, doc_local):
    """Fill the field in the distinct cell whose text contains label_sub."""
    if tbl_idx is None:
        return False
    try:
        for c in distinct(doc_local.tables[tbl_idx].rows[row]):
            if label_sub.lower() in c.text.lower():
                if not set_content_control(c, value):
                    append_val(c, value)
                return True
    except (IndexError, AttributeError) as _e:
        logger.warning('fill(%s, row=%s, %r): %s', tbl_idx, row, label_sub, _e)
    return False

def fill_spec(doc_local, tbl_idx, col_name, data, single_col):
    t = doc_local.tables[tbl_idx]
    cidx = next((i for i, c in enumerate(t.rows[0].cells)
                 if col_name.lower() in c.text.lower()), None)
    if cidx is None: return 0
    keys = sorted(data.keys(), key=len, reverse=True)
    n = 0
    for r in t.rows[1:]:
        lab = r.cells[0].text.strip()
        for k in keys:
            val = data[k].get(col_name)
            if lab.startswith(k) and val:
                if cidx < len(r.cells) and r.cells[cidx]._tc is not r.cells[0]._tc:
                    set_cell_text(r.cells[cidx], val)
                else:
                    _append(r.cells[0], val if single_col else f'[{col_name}: {val}]')
                n += 1; break
    return n


def fill_spec_by_index(doc_local, tbl_idx, col_index, col_label, data, single_col, header_override=None):
    """Like fill_spec(), but targets a column by its already-known 0-indexed position
    instead of a substring header search -- required once a column's header text isn't
    unique (the template's three identical "Limited use bmkr" cells) or was just created
    by docx_table_ops.insert_column_after and has no other way to be found yet.
    `data` is a flat {row_label: value} dict for this ONE column (unlike fill_spec()'s
    {row_label: {column_name: value}} nested shape, since the caller already knows
    exactly which single column this write targets). `header_override`, if given,
    renames the column's existing header cell to this text first (used when a
    pre-existing column is being repurposed as "<label> (1)" of a duplicate-column
    group; the (2)/(3)/... siblings already got their header text set at insertion
    time by docx_table_ops.insert_column_after, so they pass header_override=None).

    Uses docx_table_ops.row_cells() rather than python-docx's own `row.cells` --
    several data rows in the Referral Lab / Storage Samples tables wrap some
    cells in a Word content control (dropdown), which `row.cells` silently
    skips, undercounting the row and misaligning every position-based lookup.

    Every write uses docx_table_ops.strip_dropdown_and_set_text() -- if the
    target cell has an inline dropdown content control, it's DELETED and
    replaced with plain text (per direct confirmation, not appended beside
    it); otherwise it falls back to a normal plain-text overwrite. This is
    deliberately NOT limited to a fixed list of "known dropdown rows" -- e.g.
    "Sample type" turned out to be a dropdown too (a specimen-type/tube
    picklist) despite not being one of the rows originally flagged as such,
    so detecting per-cell rather than hardcoding by row label is what
    actually generalizes correctly."""
    from docx_table_ops import row_cells, strip_dropdown_and_set_text

    if tbl_idx is None or col_index is None:
        return 0
    t = doc_local.tables[tbl_idx]
    header_cells = row_cells(t.rows[0])
    if col_index >= len(header_cells):
        return 0
    if header_override:
        set_cell_text(header_cells[col_index], header_override)
    keys = sorted(data.keys(), key=len, reverse=True)
    n = 0
    for r in t.rows[1:]:
        cells = row_cells(r)
        lab = cells[0].text.strip()
        for k in keys:
            val = data.get(k)
            if lab.startswith(k) and val:
                if col_index < len(cells) and cells[col_index]._tc is not cells[0]._tc:
                    strip_dropdown_and_set_text(cells[col_index], val)
                else:
                    _append(cells[0], val if single_col else f'[{col_label}: {val}]')
                n += 1; break
    return n


def fill_shared_row(doc_local, tbl_idx, row_label_prefix, value):
    """Writes `value` ONCE into a row that's fully merged across every column in
    the table (a single shared cell for the whole row, not per-column) -- used
    for clips_nonpkpd_parser.SHARED_ROW_CONDITION_LABELS, where a per-column
    write doesn't make sense at all (there's only one cell for the entire row
    regardless of which column a file was assigned to). Finds the row by a
    startswith match on its label (same convention as fill_spec()/
    fill_spec_by_index(), robust to a multi-line label like "Ship site to
    central lab\\n         Condition: "), and appends `value` into that row's
    first cell. Returns 1 if a matching row was found and written, else 0."""
    if tbl_idx is None:
        return 0
    from docx_table_ops import row_cells

    t = doc_local.tables[tbl_idx]
    for r in t.rows[1:]:
        cells = row_cells(r)
        if cells[0].text.strip().startswith(row_label_prefix):
            _append(cells[0], value)
            return 1
    return 0


def find_table_idx(doc, keyword, min_rows=1, min_cols=1, prefer='rows'):
    """Return index of the best-matching table for a header keyword (row 0, cell 0).
    Module-level so both main()'s own table resolution and locate_specimen_tables()
    below can share one implementation instead of drifting apart."""
    candidates = []
    for _i, _t in enumerate(doc.tables):
        try:
            txt = _t.rows[0].cells[0].text.lower()
        except (IndexError, AttributeError):
            continue
        if keyword.lower() in txt:
            if len(_t.rows) >= min_rows and len(_t.columns) >= min_cols:
                candidates.append((_i, len(_t.rows), len(_t.columns)))
    if not candidates:
        return None
    if prefer == 'cols':
        return max(candidates, key=lambda x: (x[2], x[1]))[0]
    return max(candidates, key=lambda x: (x[1], x[2]))[0]


def table_after(doc, wrapper_idx):
    """Return the index of the table immediately after wrapper_idx."""
    if wrapper_idx is not None and wrapper_idx + 1 < len(doc.tables):
        return wrapper_idx + 1
    return None


def locate_specimen_tables(doc):
    """Locates the Referral Lab and two Storage Samples data tables by the same
    header-keyword + row-count heuristic main() used to run inline for T18/T20/T22 --
    extracted here (unchanged logic) so specimen_columns.py's column registry can find
    the exact same three tables without re-deriving/duplicating this heuristic a second
    time. Returns {"referral": idx, "storage_wide": idx, "storage_narrow": idx} -- any
    value may be None if that table couldn't be located."""
    t18w = find_table_idx(doc, 'REFERRAL LAB', min_rows=2)
    t18 = table_after(doc, t18w) if (t18w is not None and len(doc.tables[t18w].rows) < 4) else t18w
    if t18 is None or len(doc.tables[t18].rows) < 10:
        t18 = find_table_idx(doc, 'REFERRAL LAB', min_rows=10) or t18

    stor_candidates = []
    for si, st in enumerate(doc.tables):
        try:
            stxt = st.rows[0].cells[0].text.lower()
        except (IndexError, AttributeError):
            continue
        if 'storage samples' in stxt and len(st.rows) >= 10:
            stor_candidates.append((si, len(st.rows), len(st.columns)))
    stor_candidates.sort(key=lambda x: -x[2])  # most cols first
    t20 = stor_candidates[0][0] if len(stor_candidates) > 0 else None
    t22 = stor_candidates[1][0] if len(stor_candidates) > 1 else None
    if t20 is None:
        t20w = find_table_idx(doc, 'STORAGE SAMPLES', min_rows=2)
        t20 = table_after(doc, t20w)

    return {'referral': t18, 'storage_wide': t20, 'storage_narrow': t22}

_XML_UNSAFE = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')


def _docx_safe(value):
    """Strip control characters python-docx/lxml reject outright (confirmed:
    a bare \\x02 from PDF text extraction — likely a hyphen/ligature
    artifact — crashed the whole pipeline with no output at all). Applied at
    the write boundary so it catches bad characters regardless of which
    extraction path (edgeparse, regex fallback, design parser) produced
    them."""
    if not isinstance(value, str):
        return value
    return _XML_UNSAFE.sub('', value)


def _best_pipe_table(section_text: str, table_type: str):
    """Pick the most likely pipe-table for SoA or lab appendix from a section."""
    best_headers, best_rows, best_score = [], [], -1
    for lines in find_pipe_tables(section_text or ''):
        headers, rows = parse_pipe_table(lines)
        if not headers or not rows:
            continue

        h0 = headers[0].strip().lower() if headers else ''
        h_all = [h.strip().lower() for h in headers]
        score = len(rows)

        if table_type == 'soa':
            if any('visit' in h for h in h_all):
                score += 25
            if any(tok in h0 for tok in ('procedure', 'assessment', 'test', 'laboratory')):
                score += 12
            if len(headers) >= 4:
                score += 8
        else:
            if any('clinical' in h and 'lab' in h for h in h_all):
                score += 25
            if any(tok in h0 for tok in ('test', 'analyte', 'laboratory')):
                score += 12
            if any('comment' in h for h in h_all):
                score += 8

        if score > best_score:
            best_headers, best_rows, best_score = headers, rows, score

    return best_headers, best_rows


def insert_soa_table(doc, idx, headers, rows, footnotes, enrolled, screened, ed_rate):
    from docx.oxml import OxmlElement
    from docx.shared import Inches

    old = doc.tables[idx]._tbl

    # Free-text summary above the table replaces the old in-table "# Patients" row (this used to
    # be _prow below) -- reads more naturally than a table row of repeated counts, and matches how
    # enrollment numbers are already stated as plain text elsewhere in the template.
    _para = OxmlElement('w:p')
    _run = OxmlElement('w:r')
    _rpr = OxmlElement('w:rPr')
    _rpr.append(OxmlElement('w:b'))
    _run.append(_rpr)
    _t = OxmlElement('w:t')
    _t.text = f'Screening = {screened if screened else "?"}    Enrolled = {enrolled if enrolled else "?"}'
    _t.set(qn('xml:space'), 'preserve')
    _run.append(_t)
    _para.append(_run)
    old.addprevious(_para)

    ncols = len(headers)
    nrows = 1 + len(rows) + (1 if footnotes else 0)
    new = doc.add_table(rows=nrows, cols=ncols)
    try: new.style = 'Table Grid'
    except Exception as _e: logger.warning('insert_soa_table style: %s', _e)
    def _set(cells, values):
        for i, c in enumerate(cells):
            c.text = _docx_safe(values[i]) if i < len(values) else ''
    _set(new.rows[0].cells, headers)

    # master-table.js pads short/mismatched headers with the literal string "Column N" (1-indexed,
    # confirmed at its own source) when merging pages of differing width -- merge each such
    # placeholder cell into the last real-named header cell to its left, row 1 only, so the docx
    # never shows a bare, meaningless "Column 6"/"Column 7" label.
    _placeholder_re = re.compile(r'^Column \d+$')
    _anchor_idx = 0
    for _i in range(1, len(headers)):
        if _placeholder_re.match(headers[_i].strip()):
            try:
                # Clear the placeholder's own text first -- merge() otherwise keeps it as an
                # extra (blank) paragraph inside the merged cell, when the whole point is that
                # "Column 6" carries no real meaning and shouldn't leave any trace at all.
                new.rows[0].cells[_i].text = ''
                _merged = new.rows[0].cells[_anchor_idx].merge(new.rows[0].cells[_i])
                for _p in list(_merged.paragraphs[1:]):
                    if not _p.text.strip():
                        _p._p.getparent().remove(_p._p)
            except Exception as _e:
                logger.warning('SoA header merge (col %s): %s', _i, _e)
        else:
            _anchor_idx = _i

    for ri, _r in enumerate(rows):
        _vals = (_r + [''] * ncols)[:ncols]
        _set(new.rows[1 + ri].cells, _vals)
    if footnotes:
        _set(new.rows[-1].cells, [footnotes] + [''] * (ncols - 1))

    # Column widths: Comments reads better wider, individual visit columns narrower. These are
    # kept as *relative* sizing hints -- Word still uses each column's width as a proportion when
    # dividing up the table's overall width, even under AutoFit Window below.
    _comments_idx = next((i for i, h in enumerate(headers)
                          if h.strip().lower() in ('comments', 'comment')), None)
    for _i, _h in enumerate(headers):
        if _i == 0:
            _w = Inches(1.4)
        elif _i == _comments_idx:
            _w = Inches(1.8)
        elif is_visit_column(_h):
            _w = Inches(0.55)
        else:
            _w = Inches(0.9)
        for _row in new.rows:
            _row.cells[_i].width = _w

    # AutoFit Window: the table stretches/shrinks to always fill the page's usable width, instead
    # of being pinned to the (often much wider, for many-visit-column SoAs) sum of the fixed inch
    # widths above -- exactly Word's own "AutoFit > AutoFit Window" menu action. python-docx has no
    # high-level setter for a percentage table width (only absolute Inches/dxa), so this is set
    # directly on tblPr: w:tblW type="pct" w="5000" (5000 = 100%, in fiftieths-of-a-percent) plus
    # w:tblLayout type="autofit". Must come after new.autofit/cell-width setup above, since it
    # replaces the effect of new.autofit=False previously used here.
    _tblPr = new._tbl.tblPr
    for _tag in ('w:tblLayout', 'w:tblW'):
        _el = _tblPr.find(qn(_tag))
        if _el is not None:
            _tblPr.remove(_el)
    _tblLayout = OxmlElement('w:tblLayout')
    _tblLayout.set(qn('w:type'), 'autofit')
    _tblPr.append(_tblLayout)
    _tblW = OxmlElement('w:tblW')
    _tblW.set(qn('w:type'), 'pct')
    _tblW.set(qn('w:w'), '5000')
    _tblPr.append(_tblW)

    old.addprevious(new._tbl)
    old.getparent().remove(old)
    return len(rows)

def _parse_milestone_date(date_str):
    s = (date_str or '').strip()
    if not s: return None
    # Both full ("January") and 3-letter abbreviated ("Jan") names -- Excel's own default date
    # display format (which the "EMP" CLS Study List pull's PA/FPV/LPV dates come through
    # verbatim, via Graph's `text` cell values) is routinely abbreviated, e.g. "18-Jul-2026".
    months = {'january':1,'february':2,'march':3,'april':4,'may':5,'june':6,'july':7,
              'august':8,'september':9,'october':10,'november':11,'december':12,
              'jan':1,'feb':2,'mar':3,'apr':4,'jun':6,'jul':7,'aug':8,
              'sep':9,'sept':9,'oct':10,'nov':11,'dec':12}
    m = re.match(r'(\d{1,2})\s*[- ]\s*([A-Za-z]+)\s*[- ]\s*(\d{4})', s)
    if m:
        day, mn, yr = int(m.group(1)), m.group(2).lower(), int(m.group(3))
        return datetime.date(yr, months.get(mn, 1), day) if mn in months else None
    m = re.match(r'([A-Za-z]+)\s*[- ]\s*(\d{4})', s)
    if m:
        mn, yr = m.group(1).lower(), int(m.group(2))
        return datetime.date(yr, months.get(mn, 1), 1) if mn in months else None
    return None

def tick_checkbox(sdt, w14_ns):
    pr = sdt.find(qn('w:sdtPr')); cb = pr.find('{%s}checkbox' % w14_ns)
    chk = cb.find('{%s}checked' % w14_ns)
    from docx.oxml import OxmlElement as _Ox
    if chk is None:
        chk = _Ox('w14:checked'); cb.append(chk)
    chk.set('{%s}val' % w14_ns, '1')
    cs = cb.find('{%s}checkedState' % w14_ns)
    glyph = chr(int(cs.get('{%s}val' % w14_ns), 16)) if cs is not None else '☒'
    content = sdt.find(qn('w:sdtContent'))
    if content is not None:
        for t in content.iter(qn('w:t')): t.text = glyph

def _find_table(doc, kw):
    for t in doc.tables:
        try:
            if t.rows and kw.lower() in t.rows[0].cells[0].text.lower():
                return t
        except (IndexError, AttributeError):
            continue
    return None

def _del_table(doc, kw):
    """Delete ALL tables whose first-row first-cell text contains kw."""
    deleted = False
    # Iterate in reverse so removal doesn't shift indices
    for t in list(doc.tables):
        try:
            if t.rows and kw.lower() in t.rows[0].cells[0].text.lower():
                t._tbl.getparent().remove(t._tbl)
                deleted = True
        except (IndexError, AttributeError):
            continue
    return deleted

def _del_rows(doc, label_kw, n_after):
    """Delete rows matching label_kw (plus n_after following rows).
    If the matching row is in a table that has only those rows (i.e.,
    it's a standalone section table), delete the entire table instead."""
    for t in doc.tables:
        try:
            for i, r in enumerate(t.rows):
                if label_kw.lower() in r.cells[0].text.lower():
                    rows_to_del = list(t.rows)[i:i + 1 + n_after]
                    if len(rows_to_del) >= len(t.rows):
                        # Would delete all rows — remove the entire table
                        t._tbl.getparent().remove(t._tbl)
                    else:
                        for r2 in rows_to_del:
                            r2._tr.getparent().remove(r2._tr)
                    return True
        except (IndexError, AttributeError):
            continue
    return False


# ======================================================================
# Main entry point
# ======================================================================

def main(protocol_text: str, design_text: str,
         template_path: str, output_path: str, report_path: str,
         answers: Optional[dict] = None,
         protocol_pdf_path: str = '',
         previous_rfp_path: str = '',
         soa_include_indices: Optional[set] = None,
         soa_table_override: Optional[dict] = None,
         lab_table_override: Optional[list] = None,
         field_overrides: Optional[dict] = None,
         clips_nonpkpd_assignments: Optional[list] = None,
         previous_rfp_column_selection: Optional[dict] = None) -> Report:
    """Run the full RFP population pipeline.

    Args:
        protocol_text: Full protocol markdown/text (JSON suffix auto-stripped).
                       Used for scalar-field regex extraction (protocol number,
                       compound, phase, title, enrollment, etc.).
        design_text: Full design elements markdown/text (JSON suffix auto-stripped).
        template_path: Path to the .docx template.
        output_path: Where to save the populated .docx.
        report_path: Where to save the fill report .md.
        answers: Optional dict of user intake answers (e.g. from UI form).
        protocol_pdf_path: Path to the original protocol PDF. Used for the
                           Schedule of Activities and Lab Appendix table
                           extraction (edgeparse_extractor), which needs the
                           PDF's layout data, not flattened text. If empty,
                           falls back to text-landmark regex extraction.
        previous_rfp_path: Optional path to a previously completed Central Lab
                           RFP .docx (same template) — its Referral Lab /
                           Storage Samples tables are read directly and
                           transferred cell-by-cell.
        soa_table_override: Optional {"headers": [...], "rows": [[...]],
                           "footnotes": str} — when given, this exact
                           structure is inserted as the Schedule of Activities
                           table instead of running edgeparse/regex
                           extraction at all. Lets a caller that already has
                           its own well-structured table data (e.g. from a
                           dedicated PDF table-extraction tool) skip the
                           lossy flatten-to-text-then-re-extract round trip.
        soa_include_indices still applies on top of this override.
        lab_table_override: Optional list of [test, comment] pairs — same
                           idea as soa_table_override, but for the Lab
                           Appendix / Clinical Laboratory Tests table.
        clips_nonpkpd_assignments: Optional [{"path", "column"}, ...] -- CLIPS forms /
                           Non-PK Data Mgmt Worksheets read directly (via
                           clips_nonpkpd_parser.py) INSTEAD OF previous_rfp_path, each
                           already assigned to its target Referral/Storage column
                           ("LTS PK", "LTS Immunogenicity", "LTS DNA", "LTS Serum",
                           "LTS Plasma", "LTS RNA"). Referral/Storage has exactly one
                           source at a time -- if given, this wins outright and
                           previous_rfp_path's own Referral/Storage data is not used at
                           all (see the specimen-mgmt block below).
        previous_rfp_column_selection: Optional {table_role: [column_key, ...]}
                           (table_role one of "referral"/"storage_wide"/
                           "storage_narrow", column_key from
                           specimen_columns.list_columns()) -- which of
                           previous_rfp_path's real columns to actually use.
                           Any column NOT listed is DELETED from that table
                           in the output entirely (not just left blank) --
                           see build_specimen.py/docx_table_ops.delete_columns().
                           Only meaningful together with previous_rfp_path,
                           and only when clips_nonpkpd_assignments is NOT
                           given. If omitted (e.g. an older caller), falls
                           back to the original fixed 6-column behavior with
                           no deletion, for backward compatibility.

    Returns:
        Report dataclass with findings and coverage stats.
    """
    PT = strip_json_suffix(protocol_text)
    DT = strip_json_suffix(design_text)
    answers = answers or {}

    # ---- Design Elements Parser ---
    DESIGN_PARSER = DesignParser(DT)
    DESIGN_DATA = DESIGN_PARSER.parse()

    findings: list[Finding] = []
    def rec(field, value, source, status):
        findings.append(Finding(field=field, value=value or '', source=source, status=status))
        return value

    # field_overrides: resolved values from preparsed_schema_extractor.py (extraction_schema_v4.json) --
    # the authoritative field-resolution path going forward. Wins over this function's own
    # extraction wherever it has a non-empty value for a given field name.
    def _ov(field_name, computed):
        if field_overrides:
            v = field_overrides.get(field_name)
            if v not in (None, ''):
                return v
        return computed

    alias    = _ov('Protocol alias', protocol_number(PT))
    compound_val = _ov('Compound', compound(PT))
    title    = _ov('Protocol title', protocol_title(PT))
    phase_val = _ov('Phase', phase(PT))
    ta       = _ov('Therapeutic Area', DESIGN_DATA.flags.therapeutic_area or therapeutic_area(DT) or therapeutic_area(PT))

    immuno   = _ov('Immunogenicity testing needed',
                    'Yes' if DESIGN_DATA.flags.immunogenicity_needed else (immunogenicity(DT) or immunogenicity(PT)))
    genetics = _ov('Genetics/PGx sample collected',
                    'Yes' if DESIGN_DATA.flags.genetics_pgx_collected else (genetics_pgx(DT) or genetics_pgx(PT)))

    # Same _ov()-wins-when-non-empty precedence as every other field here — the CLS Studio app's
    # Information section can now supply a real, editable contact, replacing what used to be a
    # plain hardcoded string with no override hook at all.
    requestor_contact = _ov('Requestor contact', 'Lisa Brennan, lisa.brennan@lilly.com')
    rec('General Information — requestor contact', requestor_contact, 'profile', 'filled')
    rec('General Information — requestor phone', None, 'profile', 'review')
    today = datetime.date.today()
    def bdays(d, n):
        c = 0
        while c < n:
            d += datetime.timedelta(days=1)
            if d.weekday() < 5: c += 1
        return d
    _submitted_ov_hit = field_overrides and field_overrides.get('Date RFP submitted') not in (None, '')
    _budget_ov_hit = field_overrides and field_overrides.get('Date budget required') not in (None, '')
    submitted = _ov('Date RFP submitted', today.strftime('%d-%b-%Y'))
    budget    = _ov('Date budget required', bdays(today, 10).strftime('%d-%b-%Y'))
    rec('Date RFP submitted', submitted, 'app (Information)' if _submitted_ov_hit else 'computed (today)', 'filled' if _submitted_ov_hit else 'computed')
    rec('Date budget required', budget, 'app (Information)' if _budget_ov_hit else 'computed (+10 business days)', 'filled' if _budget_ov_hit else 'computed')

    _enroll_val = _ov('Patients enrolled (randomized)', DESIGN_DATA.enrollment.planned)
    ENROLLED = _enroll_val if _enroll_val is not None else None
    SCREEN_FAIL_RATE = (DESIGN_DATA.enrollment.screen_fail_rate
                        if DESIGN_DATA.enrollment.planned is not None
                        else screen_fail_rate(DT))
    SCREENED = _ov('Patients screened', round(ENROLLED / (1 - SCREEN_FAIL_RATE)) if ENROLLED else None)
    ED_RATE = DESIGN_DATA.enrollment.early_discontinuation_rate
    SITES_PLANNED = _ov('Sites planned', None)

    # ── Country allocation — extraction_schema_v4.json (via field_overrides) first, falling
    # back to this function's own column-4-scoped "Country Check" row scan. Schema-v4's shape
    # ({name, abbreviation, pct, status}) differs slightly from extract_country_check_row's
    # ({name, abbreviation, pct, status}) -- both now the SAME function under the hood, just
    # reached two different ways, so no shape adapter is actually needed for the fallback.
    # Deliberately NOT falling back to DESIGN_DATA.country_allocation.countries (DesignParser's
    # own broad section/keyword scan) or parse_country_allocation() -- both scan much wider
    # spans of text and pick up example/instructional text and unrelated country mentions
    # (confirmed directly: this was the actual cause of "still pulling them all in" even after
    # the column-4 fix, since that fix only ever *added* to whichever of those broad lists fed
    # it, never replaced it).
    CTRIES = []
    _country_source = 'design elements (parser)'
    _countries_ov = field_overrides.get('Country Allocation table') if field_overrides else None
    if _countries_ov:
        CTRIES = [
            {'name': c.get('name'), 'abbreviation': c.get('abbreviation') or c.get('name'),
             'pct': c.get('pct'), 'consideration': c.get('status') == 'under_consideration'}
            for c in _countries_ov if c.get('name')
        ]
        _country_source = 'app (schema v4)'
    else:
        _fallback_ctries = extract_country_check_row(DT, [])
        if _fallback_ctries:
            CTRIES = [
                {'name': c.get('name'), 'abbreviation': c.get('abbreviation') or c.get('name'),
                 'pct': c.get('pct'), 'consideration': c.get('status') == 'under_consideration'}
                for c in _fallback_ctries if c.get('name')
            ]

    COUNTRIES = [c['name'] for c in CTRIES] if CTRIES else []
    _enrolled_ov_hit = field_overrides and field_overrides.get('Patients enrolled (randomized)') not in (None, '')
    rec('Patients enrolled (randomized)', str(ENROLLED),
        'app (schema v4)' if _enrolled_ov_hit else 'design elements (parser)',
        'filled' if ENROLLED else 'review')
    _screened_ov_hit = field_overrides and field_overrides.get('Patients screened') not in (None, '')
    # The label's own screen-fail-rate percentage only reflects SCREEN_FAIL_RATE when this
    # function computed SCREENED itself -- when the value came from field_overrides, showing
    # SCREEN_FAIL_RATE next to it would be misleading (it may not be the rate that was actually
    # used upstream), so the label just states the value with no implied rate.
    _screened_label = str(SCREENED) if _screened_ov_hit else f'{SCREENED} (enrolled / (1 - {SCREEN_FAIL_RATE:.0%}))'
    rec('Patients screened', _screened_label,
        'app (schema v4)' if _screened_ov_hit else 'computed (design screen-fail rate)',
        'computed' if SCREENED else 'review')
    rec('Sites planned', str(SITES_PLANNED), 'app (CLS Study List)',
        'filled' if SITES_PLANNED else 'review')
    _in_scope_ctries = [c['name'] for c in CTRIES if not c['consideration']]
    _consider_ctries = [c['name'] for c in CTRIES if c['consideration']]
    _ctries_label = ', '.join(_in_scope_ctries)
    if _consider_ctries:
        _ctries_label += ' [consideration: ' + ', '.join(_consider_ctries) + ']'
    rec('Countries in scope', _ctries_label or '—', _country_source,
        'filled' if CTRIES else 'review')
    rec('Country Allocation table', f'{len(CTRIES)} countries ('
        f'{len(_in_scope_ctries)} in scope, {len(_consider_ctries)} consideration)',
        _country_source, 'filled' if CTRIES else 'review')

    # Pediatric determination — extraction_schema_v4.json first (field_overrides, "Yes"/"No"
    # string), then design elements, then protocol fallback
    _pediatric_ov = field_overrides.get('Pediatric population?') if field_overrides else None
    if isinstance(_pediatric_ov, str) and _pediatric_ov.strip().lower() in ('yes', 'no'):
        IS_PEDIATRIC = _pediatric_ov.strip().lower() == 'yes'
    else:
        IS_PEDIATRIC = DESIGN_DATA.flags.includes_pediatric  # Q&A table parser or section parser
    MIN_AGE = None
    if IS_PEDIATRIC is None:
        # Fallback: try to determine from protocol inclusion criteria
        _h1_age = find_heading_by_number(PT, 5)  # Section 5 = Study Population
        if not _h1_age:
            _h1_age = find_heading_by_number(PT, 1)
        _sec_age = extract_section(PT, _h1_age[0], _h1_age[1], _h1_age[2]) if _h1_age else None
        MIN_AGE = min_age(PT, section_text=_sec_age)
        if MIN_AGE is not None:
            IS_PEDIATRIC = MIN_AGE < 18
    _onc_override = str(answers.get('oncology_override', '')).strip().lower()
    if _onc_override in ('yes', 'no'):
        IS_ONCOLOGY = _onc_override == 'yes'
    else:
        IS_ONCOLOGY = is_oncology(PT + DT, ta)
    rec('Pediatric population?',
        ('undetermined' if IS_PEDIATRIC is None else 'Yes' if IS_PEDIATRIC else 'No')
        + (f' (min age {MIN_AGE})' if MIN_AGE else ''), 'protocol inclusion criteria',
        'filled' if IS_PEDIATRIC is not None else 'review')
    rec('Oncology study?', f'{"Yes" if IS_ONCOLOGY else "No"} (TA: {ta})',
        'user override' if _onc_override in ('yes', 'no') else 'design therapeutic area', 'filled')
    rec('Hepatic algorithm', 'user selects from intake (Non-oncology / Oncology / Oncology w/ ICI)', 'intake', 'filled')
    rec('Protocol alias', alias, 'protocol', 'filled' if alias else 'review')
    rec('Protocol title', title, 'protocol', 'filled' if title else 'review')
    rec('Compound', compound_val, 'protocol', 'filled' if compound_val else 'review')
    rec('Phase', phase_val, 'protocol', 'filled' if phase_val else 'review')
    rec('Therapeutic Area', ta, 'design', 'filled' if ta else 'review')
    rec('Immunogenicity testing needed', immuno, 'design', 'filled' if immuno else 'review')
    rec('Genetics/PGx sample collected', genetics, 'design', 'filled' if genetics else 'review')
    # 'Analytes -- not reported to sites' rec()+fill() moved below, after soa_rows/ana_rows are
    # both finalized -- it needs to scan them, and previously only ever recorded a hardcoded 'PK'
    # in the fill report without writing anything into the actual document at all.

    # --- Lab Appendix (Analytes table) — override > edgeparse > regex fallback ---
    ana_rows = []
    _ana_source = 'protocol'
    if lab_table_override is not None:
        for _row in lab_table_override:
            test_name = str(_row[0]).strip() if _row else ''
            comment = str(_row[1]).strip() if len(_row) > 1 else ''
            if test_name:
                ana_rows.append((test_name, comment))
        if ana_rows:
            _ana_source = 'app (structured table selection)'
    if not ana_rows and protocol_pdf_path:
        logger.info('edgeparse: extracting Lab Appendix table...')
        from edgeparse_extractor import extract_lab_appendix_table
        _ana_headers, _ana_data_rows = extract_lab_appendix_table(protocol_pdf_path)
        for _row in _ana_data_rows:
            test_name = _row[0].strip() if _row else ''
            comment = _row[1].strip() if len(_row) > 1 else ''
            if test_name:
                ana_rows.append((test_name, comment))
        if ana_rows:
            _ana_source = 'protocol (edgeparse)'
        logger.info('edgeparse: Lab Appendix -> %d analyte rows', len(ana_rows))

    # Fallback: regex-based extraction if no PDF path given or edgeparse failed
    if not ana_rows:
        _h10 = find_heading_by_number(PT, '10.2')
        if not _h10:
            _h10 = find_heading_by_number(PT, 10)
        if _h10:
            _h10_line, _h10_level, _h10_off = _h10
            _sec10 = extract_section(PT, _h10_line, _h10_level, _h10_off)
            _t_h, _t_r = _best_pipe_table(_sec10, 'lab')
            if _t_h and _t_r:
                for _row in _t_r:
                    test_name = _row[0].strip() if _row else ''
                    comment = _row[1].strip() if len(_row) > 1 else ''
                    if test_name:
                        ana_rows.append((test_name, comment))
                if ana_rows:
                    _ana_source = 'protocol (pipe-table)'
            if not ana_rows:
                ana_rows = parse_analytes_text(_sec10)

    rec('Analytes — Appendix 2 table (Test | Comments)',
        f'{len(ana_rows)} rows' if ana_rows else None,
        _ana_source,
        'filled' if ana_rows else 'review')

    _review_fields = []
    _pi_early = answers.get('penalties_incentives', False)
    if _pi_early:
        _review_fields.append(('Penalties & Incentives metrics', 'ops input not in protocol/design'))
    for fld, why in _review_fields:
        rec(fld, None, '—', 'review')

    # ---- Docx fill ---
    doc = docx.Document(str(template_path))

    # ── Table discovery ──────────────────────────────────────────────────────
    # The template has a repeating pattern: a 1-2 row title/wrapper table
    # immediately followed by the real data table.  We find the right table
    # by (a) matching the header text and (b) preferring the table with the
    # most rows / columns when multiple candidates match.
    #
    # Actual template layout (confirmed from docx inspection):
    #   0  General Information          (19 rows, 5 cols)  ← T0
    #   1  Pediatric Considerations     (4 rows)
    #   2  Oncology Considerations      (2 rows)
    #   3  Decentralized Trials         (6 rows)
    #   4  Penalties and Incentives     (3 rows)
    #   5  Metrics in-Scope             (6 rows, 4 cols)   ← T5
    #   6  Country Allocation  [title]  (2 rows, wrapper)
    #   7  Country             [data]   (5 rows, 23 cols)  ← T7
    #   8  Visit Test Schedule [title]  (2 rows, wrapper)
    #   9  SoA data table               (11 rows, 9 cols)  ← T7_SOA
    #  10  Analytes            [title]  (3 rows)           ← T8
    #  11  Reflex/Optional              (10 rows, 2 cols)  ← T9
    #  12  Hypersensitivity    [title]  (2 rows)
    #  13  Hypersensitivity    [data]   (10 rows, 2 cols)  ← T11
    #  14  Hepatic monitoring  [title]  (9 rows, 6 cols)   ← T12
    #  15  Calculations (Standard)      (4 rows)
    #  16  Anatomic Pathology  [title]  (2 rows)
    #  17  Tissue Specifications        (15 rows, 2 cols)
    #  18  Slides                       (14 rows, 2 cols)
    #  19  Referral Lab        [title]  (2 rows, wrapper)
    #  20  Referral Lab        [data]   (25 rows, 6 cols)  ← T18
    #  21  Storage Samples 1   [title]  (2 rows, wrapper)
    #  22  Storage Samples 1   [data]   (23 rows, 6 cols)  ← T20
    #  23  Storage Samples 2   [title]  (2 rows, wrapper)
    #  24  Storage Samples 2   [data]   (22 rows, 3 cols)  ← T22
    #  25  Result Management            (4 rows)           ← T25 (part)
    #  26  Bulk Supplies                (3 rows)           ← T25 (part)
    #  27  Kits                         (19 rows, 6 cols)  ← T25 (part)
    #  28  Translations                 (14 rows, 8 cols)  ← T26

    def _find_table_idx(keyword, min_rows=1, min_cols=1, prefer='rows'):
        """Return index of best matching table for a header keyword."""
        candidates = []
        for _i, _t in enumerate(doc.tables):
            try:
                txt = _t.rows[0].cells[0].text.lower()
            except (IndexError, AttributeError):
                continue
            if keyword.lower() in txt:
                if len(_t.rows) >= min_rows and len(_t.columns) >= min_cols:
                    candidates.append((_i, len(_t.rows), len(_t.columns)))
        if not candidates:
            return None
        if prefer == 'cols':
            return max(candidates, key=lambda x: (x[2], x[1]))[0]
        return max(candidates, key=lambda x: (x[1], x[2]))[0]

    def _table_after(wrapper_idx):
        """Return the index of the table immediately after wrapper_idx."""
        if wrapper_idx is not None and wrapper_idx + 1 < len(doc.tables):
            return wrapper_idx + 1
        return None

    # Find each table directly by its distinctive header text + size
    T0     = _find_table_idx('General Information', min_rows=5) or 0
    T5     = _find_table_idx('Metrics in-Scope', min_rows=3)
    # Country: the title wrapper is 2 rows; data table follows immediately
    _t7w   = _find_table_idx('Country Allocation', min_rows=2)
    T7     = _table_after(_t7w) if (_t7w is not None and len(doc.tables[_t7w].rows) < 4) else _t7w
    # Prefer the data table (≥5 rows, ≥5 cols) over the wrapper
    if T7 is not None and len(doc.tables[T7].rows) < 5:
        T7 = _find_table_idx('Country', min_rows=5, min_cols=5) or T7
    # SoA: title wrapper then data table
    _t7sw  = _find_table_idx('Visit Test Schedule', min_rows=2)
    T7_SOA = _table_after(_t7sw) if (_t7sw is not None and len(doc.tables[_t7sw].rows) < 4) else _t7sw
    if T7_SOA is None:
        T7_SOA = _find_table_idx('Schedule of Activities', min_rows=3)
    T8     = _find_table_idx('Analytes (Clinical Laboratory Tests)', min_rows=2) or 10
    T9     = _find_table_idx('Reflex/Optional', min_rows=4) or 11
    # Hypersensitivity: title is 2 rows; data table follows
    _t11w  = _find_table_idx('Hypersensitivity', min_rows=2)
    T11    = _table_after(_t11w) if (_t11w is not None and len(doc.tables[_t11w].rows) < 4) else _t11w
    if T11 is None:
        T11 = _find_table_idx('Visit Name', min_rows=4)
    # Hepatic: the main table has the header and data combined (≥6 rows)
    T12    = _find_table_idx('Hepatic monitoring', min_rows=5)
    if T12 is None:
        T12 = _find_table_idx('Hepatic', min_rows=4)
    # Referral Lab + Storage Samples tables: shared heuristic, extracted to module
    # level as locate_specimen_tables() so specimen_columns.py's column registry
    # (used by the CLIPS/Non-PKPD column dropdown) can find the same three tables.
    _specimen_tables = locate_specimen_tables(doc)
    T18 = _specimen_tables['referral']
    T20 = _specimen_tables['storage_wide']
    T22 = _specimen_tables['storage_narrow']
    # Result Management / Bulk Supplies / Kits — any of these identifies T25 area
    T25 = (_find_table_idx('Result Management', min_rows=2)
           or _find_table_idx('Bulk Supplies', min_rows=2)
           or _find_table_idx('Kits', min_rows=4))
    T26 = _find_table_idx('Translations', min_rows=4, min_cols=4) or _find_table_idx('Translations', min_rows=3)

    # Warn on any table we couldn't locate
    for _tname, _tidx in [('T0',T0),('T5',T5),('T7',T7),('T7_SOA',T7_SOA),
                           ('T8',T8),('T9',T9),('T11',T11),('T12',T12),
                           ('T18',T18),('T20',T20),('T22',T22),('T25',T25),('T26',T26)]:
        if _tidx is None:
            logger.warning('Table %s not found in template — fills for this section will be skipped', _tname)

    # Log resolved indices for diagnostics
    logger.info('Table indices resolved: T0=%s T5=%s T7=%s T7_SOA=%s T8=%s T9=%s '
                'T11=%s T12=%s T18=%s T20=%s T22=%s T25=%s T26=%s',
                T0, T5, T7, T7_SOA, T8, T9, T11, T12, T18, T20, T22, T25, T26)


    fill(T0, 1, 'requestor contact information', requestor_contact, doc)
    fill(T0, 2, 'Date RFP submitted', submitted, doc)
    fill(T0, 2, 'Date budget required', budget, doc)
    _other_studies_ans = answers.get('rfp_for_other_studies')
    if _other_studies_ans is not None:
        _other_studies_val = 'Yes' if _other_studies_ans else 'No'
        fill(T0, 3, 'other studies in this program', _other_studies_val, doc)
        rec('RFP for other studies in program?', _other_studies_val, 'app (Information)', 'filled')
    fill(T0, 8, 'Protocol alias', alias or REVIEW('Protocol alias'), doc)
    fill(T0, 8, 'Protocol title', title or REVIEW('Protocol title'), doc)
    fill(T0, 9, 'Compound', compound_val or REVIEW('Compound'), doc)
    fill(T0, 9, 'Phase', phase_val or REVIEW('Phase'), doc)
    fill(T0, 9, 'Therapeutic Area', ta or REVIEW('Therapeutic Area'), doc)

    _milestones_map = DESIGN_DATA.timeline
    _milestones = {
        'Protocol Approval': _milestones_map.protocol_approval,
        'FPV': _milestones_map.fpv,
        'LPV': _milestones_map.lpv,
        'FPET': _milestones_map.fpet,
        'LPET': _milestones_map.lpet,
        'Protocol Content Lock': _milestones_map.protocol_content_lock,
        'Design Element Alignment': _milestones_map.design_element_alignment,
    }
    # extraction_schema_v4.json's own milestone resolution (field_overrides) wins per-key over
    # this function's own DESIGN_DATA.timeline parse wherever it has a value.
    _milestones_ov = field_overrides.get('Trial Milestones (PA/FPV/LPV/FPET/LPET)') if field_overrides else None
    if _milestones_ov:
        _key_map = {'Protocol Approval': 'protocol_approval', 'FPV': 'fpv', 'LPV': 'lpv',
                    'FPET': 'fpet', 'LPET': 'lpet'}
        for _label, _ov_key in _key_map.items():
            _v = _milestones_ov.get(_ov_key)
            if _v not in (None, ''):
                _milestones[_label] = _v
    _pa_date = _milestones.get('Protocol Approval')
    _fpv_date = _milestones.get('FPV')
    _lpv_date = _milestones.get('LPV')
    # Tracked as a Finding (not just written via fill()) so it's visible in
    # the report/UI with a correct source label — previously these dates
    # were written straight into the template with no Finding at all,
    # making the sourcing invisible/unauditable.
    _milestone_label = ', '.join(f'{k}: {v}' for k, v in _milestones.items() if v)
    rec('Trial Milestones (PA/FPV/LPV/FPET/LPET)', _milestone_label or '—', 'design elements (parser)',
        'filled' if _milestone_label else 'review')
    if _pa_date: fill(T0, 11, 'Protocol Approval (PA) date', _pa_date, doc)
    if _fpv_date: fill(T0, 11, 'Planned First Patient Visit (FPV) date', _fpv_date, doc)
    if _fpv_date: fill(T0, 11, 'Planned First Patient Visit (FPV) date', _fpv_date, doc)
    if _lpv_date: fill(T0, 11, 'Planned Last Patient Visit (LPV) date', _lpv_date, doc)
    _fpv_d = _parse_milestone_date(_fpv_date)
    _lpv_d = _parse_milestone_date(_lpv_date)
    if _fpv_d and _lpv_d:
        # Protocol duration (FPV-DBL) = months(FPV -> LPV), plus a flat +1 month covering the
        # ~4-week data-cleaning/database-lock tail after last patient visit -- reported as a
        # single rounded month count (e.g. "15 months"), not "14 months + 4 weeks".
        _months = (_lpv_d.year - _fpv_d.year) * 12 + (_lpv_d.month - _fpv_d.month) + 1 + 1
        fill(T0, 12, 'Protocol duration (FPV-DBL)', f'{_months} months', doc)
    rec('Protocol duration (FPV-DBL)', f'{_months} months' if (_fpv_d and _lpv_d) else None,
        'computed (FPV to LPV + 1 month)', 'computed' if (_fpv_d and _lpv_d) else 'review')
    fill(T0, 13, 'Country where initial FPV planned', 'US (REVIEW)', doc)
    rec('Country where initial FPV planned', 'US (REVIEW)', 'assumed default (not derived)', 'review')
    if _fpv_d:
        _siv_d = _fpv_d - datetime.timedelta(days=14)
        _siv_str = f'{_siv_d.day}-{_siv_d.strftime("%B")}-{_siv_d.year}'
        fill(T0, 14, 'Initial SIV date', _siv_str, doc)
    rec('Initial SIV date', _siv_str if _fpv_d else None,
        'computed (FPV - 14 days)', 'computed' if _fpv_d else 'review')

    fill(T0, 18, 'Type of central lab investigator training support required', 'Training slides only', doc)
    rec('Type of central lab investigator training support required', 'Training slides only',
        'fixed default', 'filled')
    if _fpv_d:
        _train_d = _fpv_d - datetime.timedelta(weeks=3)
        _train_str = f'{_train_d.day}-{_train_d.strftime("%B")}-{_train_d.year}'
        fill(T0, 18, 'Date training slides required', _train_str, doc)
    rec('Date training slides required', _train_str if _fpv_d else None,
        'computed (FPV - 3 weeks)', 'computed' if _fpv_d else 'review')

    # Blood Volume per Visit Summary required (Required for China and Ped studies): Yes if either
    # condition holds, No only when both are No/undetermined-false.
    _china_in_scope = any('china' in (c or '').lower() for c in COUNTRIES)
    _blood_vol_val = 'Yes' if (IS_PEDIATRIC or _china_in_scope) else 'No'
    fill(T0, 5, 'Blood Volume per Visit Summary', _blood_vol_val, doc)
    rec('Blood Volume per Visit Summary required', _blood_vol_val,
        'computed (pediatric or China in scope)', 'computed')

    if T8 is not None:
        # Row 1 ("...NOT be reported to the sites...") is filled later, after soa_rows/ana_rows
        # are available -- see the "Analytes -- not reported to sites" block below, which
        # replaced this old hardcoded-'PK' call.
        try:
            for c in distinct(doc.tables[T8].rows[2]):
                if 'analyte listing' in c.text.lower():
                    append_val(c, '(reconstructed as the table below, from protocol Appendix 2)'); break
        except (IndexError, AttributeError):
            pass

    # --- Schedule of Activities (override > edgeparse > regex fallback) ---
    soa_visits, soa_rows, soa_footnotes = [], [], ''
    _soa_source = 'protocol (regex)'

    if soa_table_override is not None:
        soa_visits = list(soa_table_override.get('headers') or [])
        soa_rows = [list(r) for r in (soa_table_override.get('rows') or [])]
        soa_footnotes = soa_table_override.get('footnotes') or ''
        if soa_visits and soa_rows:
            _soa_source = 'app (structured table selection)'

    if not soa_rows and protocol_pdf_path:
        logger.info('edgeparse: extracting Schedule of Activities table...')
        from edgeparse_extractor import extract_soa_table
        soa_visits, soa_rows = extract_soa_table(protocol_pdf_path)
        if soa_visits and soa_rows:
            _soa_source = 'protocol (edgeparse)'
            logger.info('edgeparse: SoA -> %d rows x %d cols', len(soa_rows), len(soa_visits))

    # Fallback: regex-based extraction if no PDF path given or edgeparse failed
    if not soa_rows:
        _h_soa = find_heading_by_number(PT, '1.3')
        if not _h_soa:
            _h_soa = find_heading_by_number(PT, 1)
        if _h_soa:
            _h_soa_line, _h_soa_level, _h_soa_off = _h_soa
            _sec_soa = extract_section(PT, _h_soa_line, _h_soa_level, _h_soa_off)
            _pipe_h, _pipe_r = _best_pipe_table(_sec_soa, 'soa')
            if _pipe_h and _pipe_r:
                _filtered = []
                for _r in _pipe_r:
                    _label = _r[0] if _r else ''
                    _cells = _r[1:] if len(_r) > 1 else []
                    if is_category_header(_label, _cells):
                        if not is_non_lab(_label):
                            _filtered.append(_r)
                    elif not is_non_lab(_label):
                        _filtered.append(_r)
                if _filtered:
                    soa_visits = _pipe_h
                    soa_rows = _filtered
                    soa_footnotes = ''
                    _soa_source = 'protocol (pipe-table)'
            if not soa_rows:
                _raw_h, _raw_r, _raw_f = parse_soa_text(_sec_soa)
                if _raw_r:
                    _filtered = []
                    for _r in _raw_r:
                        _label = _r[0] if _r else ''
                        _cells = _r[1:] if len(_r) > 1 else []
                        if is_category_header(_label, _cells):
                            if not is_non_lab(_label):
                                _filtered.append(_r)
                        elif not is_non_lab(_label):
                            _filtered.append(_r)
                    soa_visits = _raw_h
                    soa_rows = _filtered
                    soa_footnotes = _raw_f

    # Interactive row selection (rfp-ui "crop tool"): filter by position,
    # not label text — duplicate labels are real (e.g. unmerged fragments
    # like "Exploratory biomarker" vs "Exploratory biomarker samples" can
    # coexist), so filtering by string would ambiguously affect both.
    if soa_include_indices is not None:
        soa_rows = [r for i, r in enumerate(soa_rows) if i in soa_include_indices]

    if soa_visits and soa_rows:
        soa_n = insert_soa_table(doc, T7_SOA, soa_visits, soa_rows, soa_footnotes, ENROLLED, SCREENED, ED_RATE)
        _col_str = f'{len(soa_visits)} cols' if not soa_footnotes else f'{len(soa_visits)} cols + footnotes'
        rec('Schedule of Activities matrix', f'{soa_n} lab rows x {_col_str}',
            _soa_source, 'filled')
    else:
        rec('Schedule of Activities matrix', 'not available',
            'protocol', 'review')

    # Analytes -- not reported to sites: central lab commonly doesn't report certain specialty
    # assay results directly to sites for patient management. Scans the *whole* protocol (this
    # can legitimately be called out in section 10's own Analytes listing, not just earlier
    # sections) plus this document's own SoA/Lab Appendix rows -- not just a hardcoded 'PK'.
    _nr_blob = (PT + ' '
                + ' '.join((_r[0] if _r else '') for _r in (soa_rows or []))
                + ' ' + ' '.join(_t for _t, _ in (ana_rows or []))).lower()
    _not_reported = []
    if re.search(r'\bpk\b|pharmacokinetic', _nr_blob):
        _not_reported.append('PK')
    if re.search(r'\bly ?\d{3,}\b', _nr_blob):
        _not_reported.append(compound_val if compound_val else 'LY compound assay')
    if re.search(r'\bada\b|anti[- ]drug antibod', _nr_blob):
        _not_reported.append('ADA')
    if str(immuno).strip().lower().startswith('yes'):
        _not_reported.append('Immunogenicity')
    if 'biomarker' in _nr_blob:
        _not_reported.append('Biomarkers')
    _not_reported_val = ', '.join(dict.fromkeys(_not_reported)) if _not_reported else None
    if _not_reported_val:
        _t8 = doc.tables[T8]
        for _row in _t8.rows:
            _dc = distinct(_row)
            if 'will not be reported' in _dc[0].text.lower():
                if len(_dc) > 1:
                    set_cell_text(_dc[1], _not_reported_val)
                else:
                    _append(_dc[0], _not_reported_val)
                break
    rec('Analytes — not reported to sites', _not_reported_val or 'none identified',
        'protocol + SoA/Lab Appendix scan', 'filled' if _not_reported_val else 'review')

    # Referral Lab Samples / Storage Samples have exactly one source at a time -- Previous RFP
    # (reverse-parsed via build_specimen.py) OR CLIPS forms/Non-PK Data Mgmt Worksheets (read
    # directly via clips_nonpkpd_parser.py), never both (the app's own UI already enforces this
    # by disabling each attachment path while the other has data; this is the backstop for any
    # other caller of this pipeline). CLIPS/Non-PKPD wins outright when both are somehow given,
    # rather than blending per-column, so the fill report always names a single, unambiguous
    # source instead of a silent merge.
    _clips_unmapped: list[str] = []
    if clips_nonpkpd_assignments:
        from clips_nonpkpd_parser import build_from_assignments
        try:
            _writes, _insertions, _shared_writes = build_from_assignments(clips_nonpkpd_assignments)
        except Exception as _e:
            logger.warning('clips_nonpkpd_parser failed: %s', _e)
            _writes, _insertions, _shared_writes = [], [], []
        _clips_unmapped = [a['path'] for a in clips_nonpkpd_assignments if not a.get('column')]

        _table_by_role = {'referral': T18, 'storage_wide': T20, 'storage_narrow': T22}
        from docx_table_ops import insert_column_after
        for _ins in _insertions:
            _tidx = _table_by_role.get(_ins['table_role'])
            if _tidx is None:
                continue
            try:
                insert_column_after(doc.tables[_tidx], _ins['after_col_index'], _ins['header_text'])
            except Exception as _e:
                logger.warning('insert_column_after failed for %s: %s', _ins, _e)

        for _sw in _shared_writes:
            _tidx = _table_by_role.get(_sw['table_role'])
            try:
                fill_shared_row(doc, _tidx, _sw['row_label'], _sw['value'])
            except Exception as _e:
                logger.warning('fill_shared_row failed for %s: %s', _sw, _e)

        nref = 0
        nsto = 0
        for _w in _writes:
            _tidx = _table_by_role.get(_w['table_role'])
            _n = fill_spec_by_index(
                doc, _tidx, _w['col_index'],
                _w.get('header_override') or _w['base_label'],
                _w['row_data'],
                single_col=(_w['base_label'] == 'LTS PK'),
                header_override=_w.get('header_override'),
            )
            if _w['table_role'] == 'referral':
                nref += _n
            else:
                nsto += _n
    else:
        from build_specimen import build as build_specimen_preview
        try:
            _preview = build_specimen_preview(previous_rfp_path) if previous_rfp_path else {}
        except Exception as _e:
            logger.warning('build_specimen failed: %s', _e)
            _preview = {}

        _table_by_role = {'referral': T18, 'storage_wide': T20, 'storage_narrow': T22}
        nref = 0
        nsto = 0

        if previous_rfp_column_selection is not None:
            # New flow: delete every column NOT selected, write only the selected ones.
            from docx_table_ops import delete_columns
            for _table_role, _table_data in _preview.items():
                _tidx = _table_by_role.get(_table_role)
                if _tidx is None:
                    continue
                _selected_keys = set(previous_rfp_column_selection.get(_table_role) or [])
                _columns = _table_data['columns']
                _selected_cols = [c for c in _columns if c['key'] in _selected_keys]
                _unselected_indices = [c['col_index'] for c in _columns if c['key'] not in _selected_keys]
                try:
                    delete_columns(doc.tables[_tidx], _unselected_indices)
                except Exception as _e:
                    logger.warning('delete_columns failed for %s: %s', _table_role, _e)
                    continue
                for _col in sorted(_selected_cols, key=lambda c: c['col_index']):
                    _shift = sum(1 for u in _unselected_indices if u < _col['col_index'])
                    _row_data = {
                        _label: _vals[_col['key']]
                        for _label, _vals in _table_data['rows'].items()
                        if _col['key'] in _vals
                    }
                    _n = fill_spec_by_index(
                        doc, _tidx, _col['col_index'] - _shift, _col['base_label'], _row_data,
                        single_col=(_col['base_label'] == 'LTS PK'),
                    )
                    if _table_role == 'referral':
                        nref += _n
                    else:
                        nsto += _n
        else:
            # Backward-compatible fallback (no selection given): the original fixed
            # 6-column behavior, no deletion -- reshape the preview's {key: value} rows
            # back into fill_spec()'s {column_name: value} shape.
            def _to_old_shape(table_role):
                table_data = _preview.get(table_role) or {}
                key_to_label = {c['key']: c['base_label'] for c in table_data.get('columns', [])}
                return {
                    row_label: {key_to_label[k]: v for k, v in vals.items() if k in key_to_label}
                    for row_label, vals in table_data.get('rows', {}).items()
                }

            referral = _to_old_shape('referral')
            storage = _to_old_shape('storage_wide')
            for _label, _cols in _to_old_shape('storage_narrow').items():
                storage.setdefault(_label, {}).update(_cols)

            nref = fill_spec(doc, T18, 'LTS PK', referral, single_col=True)
            nref += fill_spec(doc, T18, 'LTS Immunogenicity', referral, single_col=False)
            nsto = sum(fill_spec(doc, T20, c, storage, single_col=False) for c in ('LTS DNA', 'LTS Serum', 'LTS Plasma'))
            nsto += fill_spec(doc, T22, 'LTS RNA', storage, single_col=False)

    if clips_nonpkpd_assignments:
        _specimen_source = f"{len([a for a in clips_nonpkpd_assignments if a.get('column')])} CLIPS/Non-PKPD file(s)"
    elif previous_rfp_path:
        _specimen_source = Path(previous_rfp_path).name
    else:
        _specimen_source = 'no previous RFP or CLIPS/Non-PKPD files provided'
    rec('Specimen Mgmt — Referral Lab Samples (LTS PK / LTS Immunogenicity)', f'{nref} rows', _specimen_source,
        'filled' if nref else 'review')
    rec('Specimen Mgmt — Storage Samples (DNA/Serum/Plasma/RNA)', f'{nsto} cells', _specimen_source,
        'filled' if nsto else 'review')
    if _clips_unmapped:
        rec('Specimen Mgmt — unassigned CLIPS/Non-PKPD file(s)', '; '.join(_clips_unmapped),
            'app (needs manual column assignment)', 'review')

    # Specimens section (Storage & Shipping / Kit, Results & Data Transfer) -- report-only for
    # now (rec(), no fill()): confirmed via a full-document text scan that no template row
    # matches "storage condition"/"kit type"/"shipping frequency"/"data transfer format" anywhere,
    # so there's nowhere in the docx to write these yet. Tracked here so they're visible in the
    # fill report rather than silently dropped, and easy to wire to a real cell later if one is
    # identified.
    for _label, _ans_key in (
        ('Specimen storage conditions', 'storage_conditions'),
        ('Specimen kit type', 'kit_type'),
        ('Specimen shipping frequency', 'shipping_frequency'),
        ('Specimen data transfer format', 'data_transfer_format'),
    ):
        _val = str(answers.get(_ans_key, '') or '').strip()
        rec(_label, _val or None, 'app (Specimens)', 'filled' if _val else 'review')

    # ---- Logic Rules ---
    W14 = 'http://schemas.microsoft.com/office/word/2010/wordml'
    DEFAULTS = {
        'Database Modifications': '5 (default)', 'Sample metadata': 'Monthly (default)',
        'Header reconciliation': 'Monthly (default)', 'Data transfer to LabsConnect': 'Weekly (default)',
        'Kit overage': '50% (default)', 'Expedited production orders': '10% (default)',
        'weekend pick-up': '10% (default)',
        'Transportation assumptions': '80% primary cities, 10% secondary cities, 10% tertiary cities (default)',
    }
    nd = 0
    if T25 is not None:
        for r in doc.tables[T25].rows:
            dc = distinct(r)
            for i, c in enumerate(dc):
                for lab, val in DEFAULTS.items():
                    if lab.lower() in c.text.lower():
                        if set_content_control(c, val):
                            nd += 1
                        elif i + 1 < len(dc) and set_content_control(dc[i+1], val):
                            nd += 1
                        break
    for sdt in doc.element.body.iter(qn('w:sdt')):
        pr = sdt.find(qn('w:sdtPr'))
        dd = pr.find(qn('w:dropDownList')) if pr is not None else None
        if dd is None: continue
        deflt = next((li for li in dd.findall(qn('w:listItem'))
                      if 'default' in (li.get(qn('w:displayText')) or '').lower()), None)
        content = sdt.find(qn('w:sdtContent'))
        if deflt is not None and content is not None:
            ts = content.findall('.//' + qn('w:t'))
            if ts:
                ts[0].text = deflt.get(qn('w:displayText')) or deflt.get(qn('w:value'))
                for x in ts[1:]: x.text = ''
                nd += 1
    rec('Rule 1 — standard defaults selected', f'{nd} fields', 'standard defaults', 'filled')

    # _working_ctries must be defined before the country-table block so that
    # Rule 3 (translations) can use it even when the country table is absent.
    _working_ctries = list(CTRIES)
    _country_src = _country_source

    try:
        ca = doc.tables[T7]
    except (IndexError, KeyError):
        ca = None
    if ca is None or len(ca.rows) < 5:
        logger.warning('Country table (T7=%s) has %s rows, skipping country allocation', T7, len(ca.rows) if ca else 0)
        ca = None
    if ca is not None:
        hdr_cells = distinct(ca.rows[0])
        screened_cells = distinct(ca.rows[2])
        rand_cells = distinct(ca.rows[3])
        sites_cells = distinct(ca.rows[1])
        fpv_cells = distinct(ca.rows[4])
        _manual_raw = answers.get('country_allocation', '').strip()
        if _manual_raw:
            parsed_manual = parse_manual_countries(_manual_raw)
            if parsed_manual:
                _working_ctries = parsed_manual
                _country_src = 'manual intake'

        for j, ctr in enumerate(_working_ctries):
            col = 1 + j
            if col >= len(hdr_cells) - 1:
                break
            name = ctr['name']
            if ctr['consideration']:
                name += ' (consideration)'
            set_cell_text(hdr_cells[col], name)
            if ctr['pct'] is not None and ENROLLED:
                n_sc = round(SCREENED * ctr['pct']) if SCREENED else None
                n_en = round(ENROLLED * ctr['pct'])
                if n_sc is not None and col < len(screened_cells) - 1:
                    set_cell_text(screened_cells[col], str(n_sc))
                if col < len(rand_cells) - 1:
                    set_cell_text(rand_cells[col], str(n_en))
            if col < len(sites_cells) - 1:
                if not sites_cells[col].text.strip():
                    set_cell_text(sites_cells[col], '—')
        if SCREENED: set_cell_text(screened_cells[-1], str(SCREENED))
        if ENROLLED: set_cell_text(rand_cells[-1], str(ENROLLED))
        if SITES_PLANNED: set_cell_text(sites_cells[-1], str(SITES_PLANNED))

        for tnode in ca._tbl.iter(qn('w:t')):
            if tnode.text and 'country here' in tnode.text.lower():
                tnode.text = ''

        _n_in_scope = sum(1 for c in _working_ctries if not c['consideration'])
        _n_consider = sum(1 for c in _working_ctries if c['consideration'])
        _n_with_pct = sum(1 for c in _working_ctries if c['pct'] is not None)
        rec('Country Allocation table',
            f'{len(_working_ctries)} countries ({_n_in_scope} in scope'
            + (f', {_n_consider} consideration' if _n_consider else '')
            + f'); per-country breakdown: {_n_with_pct} with pct targets',
            _country_src, 'filled')

    N_VISITS = sum(1 for _v in (soa_visits or [])[1:] if is_visit_column(_v))
    pv = round(ENROLLED * N_VISITS * 0.02) if ENROLLED else None
    if pv:
        for ti in (T11, T12):
            if ti is None:
                continue
            try:
                for r in doc.tables[ti].rows:
                    c0 = r.cells[0].text.lower()
                    if 'patient' in c0 and 'visit' in c0:
                        for c in distinct(r)[1:]:
                            set_cell_text(c, str(pv))
                        break
            except (IndexError, AttributeError):
                pass
    rec('Rule 2 — # patient visits (hypersensitivity & hepatic)',
        f'{pv}  ({ENROLLED} enrolled x {N_VISITS} visits x 2%)', 'computed', 'computed')

    tick_langs = set()
    for ctr in _working_ctries:
        tick_langs.update(COUNTRY_LANG.get(ctr['abbreviation'], []))

    n_tick = 0
    if T26 is not None:
        t26 = doc.tables[T26]
        ALL_LANGS = set()
        for el in t26._tbl.iter(qn('w:t')):
            s = (el.text or '').strip()
            if len(s) > 2 and any(ch.isalpha() for ch in s) and s not in ('Translations',):
                ALL_LANGS.add(s)

        pending = None
        for el in t26._tbl.iter():
            tag = el.tag.split('}')[-1]
            if tag == 'sdt':
                pr = el.find(qn('w:sdtPr'))
                if pr is not None and pr.find('{%s}checkbox' % W14) is not None:
                    pending = el
            elif tag == 't' and (el.text or '').strip() in ALL_LANGS:
                if pending is not None:
                    if (el.text or '').strip() in tick_langs:
                        tick_checkbox(pending, W14); n_tick += 1
                    pending = None
    rec('Rule 3 — translations ticked', f'{n_tick}: {sorted(tick_langs)}', 'countries -> native language', 'filled')

    # Reflex/optional testing: defaults are the user's own fixed standard values (no longer gated
    # on a protocol-text keyword search) -- an analyte's row only stays on the form if it's
    # actually part of *this* study's own SoA or Lab Appendix; otherwise the whole row is removed
    # rather than shown with a "No", since asking about reflex testing for something not even in
    # the panel doesn't apply.
    REFLEX_DEFAULTS = {
        'pregnancy':  '50% of pts',
        'ck-mb':      'Yes - 1% of CK visits',
        'uds':        'Yes - assume 2%',
        'hbv dna':    'Yes - 1%',
        'hcv rna':    'Yes - 1%',
        'hiv':        'Yes - 1%',
        'fsh':        'Yes - 50%',
    }
    REFLEX_PRESENCE_KEYWORDS = {
        'pregnancy':  ['pregnan'],
        # NOT 'creatine kinase' -- that's the base CK test, present in nearly every Lab Appendix
        # regardless of whether the CK-MB *reflex* sub-test is actually part of this study, which
        # made this row's presence check a false positive almost every time.
        'ck-mb':      ['ck-mb', 'ck mb'],
        'uds':        ['uds', 'urine drug'],
        'hbv dna':    ['hbv'],
        'hcv rna':    ['hcv'],
        'hiv':        ['hiv'],
        'fsh':        ['fsh', 'follicle'],
    }
    if T9 is not None:
        t9 = doc.tables[T9]
        _soa_lab_blob = (' '.join((_r[0] if _r else '') for _r in (soa_rows or []))
                         + ' ' + ' '.join(_t for _t, _ in (ana_rows or []))).lower()
        _rows_to_drop = []
        for r in t9.rows:
            dc = distinct(r)
            lab = dc[0].text.strip().lower()
            if 'other' in lab and len(lab) < 12:
                note = ('MMA reflexed if B12 below central lab reference range; '
                        'ANA reflex to titre and pattern if positive')
                for sdt in t9._tbl.iter(qn('w:sdt')):
                    content = sdt.find(qn('w:sdtContent'))
                    if content is None: continue
                    txt = ''.join((x.text or '') for x in content.iter(qn('w:t')))
                    if 'other reflex' in txt.lower():
                        ts = content.findall('.//' + qn('w:t'))
                        ts[0].text = note
                        for x in ts[1:]: x.text = ''
                        for col in content.iter(qn('w:color')): col.set(qn('w:val'), '000000')
                        pr = sdt.find(qn('w:sdtPr'))
                        if pr is not None:
                            for tag in ('w:placeholder', 'w:showingPlcHdr'):
                                el = pr.find(qn(tag))
                                if el is not None: pr.remove(el)
                        break
                continue
            for key, deflt in REFLEX_DEFAULTS.items():
                if key in lab:
                    present = any(kw in _soa_lab_blob for kw in REFLEX_PRESENCE_KEYWORDS[key])
                    if present:
                        if len(dc) > 1:
                            set_cell_text(dc[1], deflt)
                    else:
                        _rows_to_drop.append(r)
                    break
        for r in _rows_to_drop:
            r._tr.getparent().remove(r._tr)
    rec('Rule 4 — reflex/optional testing', 'defaults applied; rows removed for analytes absent from SoA/Lab Appendix', 'protocol SoA + appendix', 'filled')

    from docx.oxml import OxmlElement
    atbl = doc.add_table(rows=1, cols=2)
    try: atbl.style = 'Table Grid'
    except Exception as _e: logger.warning('analyte table style: %s', _e)
    atbl.rows[0].cells[0].text = 'Clinical Laboratory Tests'
    atbl.rows[0].cells[1].text = 'Comments'
    for test, comment in (ana_rows or []):
        rc = atbl.add_row().cells
        rc[0].text = _docx_safe(test)
        rc[1].text = _docx_safe(comment.replace('I f ', 'If '))
    if T8 is not None:
        _t8 = doc.tables[T8]._tbl
        _spacer = OxmlElement('w:p')
        _t8.addnext(_spacer)
        _spacer.addnext(atbl._tbl)

    # ---- Conditional sections ---
    ANSWERS_DICT = {
        'pediatrics': IS_PEDIATRIC,
        'decentralized': answers.get('decentralized', False),
        'penalties_incentives': answers.get('penalties_incentives', False),
        'anatomic_pathology': answers.get('anatomic_pathology', False),
        'hepatic_calc': answers.get('hepatic_calc', 'Non-oncology'),
    }
    # Each intake toggle always resolves to a real value (they default False/'Non-oncology' when
    # unanswered, never null) -- tracked individually here, in addition to the aggregate
    # 'Conditional sections (intake)' entry below, so the toggle's own value is auditable in the
    # fill report, not just its downstream effect on which template sections got removed.
    rec('Decentralized trial elements?', 'Yes' if ANSWERS_DICT['decentralized'] else 'No',
        'user intake (Study Details panel)', 'filled')
    rec('Penalties & Incentives metrics apply?', 'Yes' if ANSWERS_DICT['penalties_incentives'] else 'No',
        'user intake (Study Details panel)', 'filled')
    rec('Anatomic Pathology / histology samples?', 'Yes' if ANSWERS_DICT['anatomic_pathology'] else 'No',
        'user intake (Study Details panel)', 'filled')

    def _ap_fill(table_kw, label_kw, value):
        t = _find_table(doc, table_kw)
        if t is None: return False
        for r in t.rows:
            if label_kw.lower() in r.cells[0].text.lower():
                dcs = distinct(r)
                if len(dcs) > 1:
                    # A real, separate answer cell (or a content control within it) -- safe to
                    # overwrite outright, same as everywhere else in this file.
                    set_cell_text(dcs[-1], value)
                else:
                    # No separate answer cell and no content control (confirmed directly against
                    # several of these template rows) -- dcs[0] *is* the label. set_cell_text
                    # would silently replace the question text with the answer, destroying it;
                    # append instead, exactly like the T25 standard-defaults loop already does
                    # for the same "label-only cell" shape.
                    _append(dcs[0], value)
                return True
        return False

    # (table keyword, row-label keyword, oncology_biopsy_extractor.py rfp_mapping key) -- matched
    # against the real template's Tissue Specifications (15 rows) and Slides (14 rows) tables,
    # confirmed by direct docx inspection. The extractor's own keys already read almost verbatim
    # off those two tables, since it was written against this exact template.
    _AP_FIELD_MAP = [
        ('Anatomic Pathology', 'summarize AP samples',
         'Anatomic Pathology / Histology - brief summary'),
        ('Tissue Specifications', 'archived or fresh',
         'Will submitted samples be archived or fresh? If combo, specify percentage of each.'),
        ('Tissue Specifications', 'tumor assessment required on an h&e',
         'Is % tumor assessment required on an H&E slide for block/slides submitted?'),
        ('Tissue Specifications', 'sample inspection required on blocks',
         'Is sample inspection required on blocks/slides received from site before storage?'),
        ('Tissue Specifications', 'h&e level of complexity', 'H&E level of complexity'),
        ('Tissue Specifications', 'is h&e required on each block',
         'If multiple blocks are received, is H&E required on each block?'),
        ('Tissue Specifications', 'cut upon receipt or upon request',
         'Should slides be cut upon receipt or upon request?'),
        ('Tissue Specifications', 'tissue curls required', 'Are tissue curls required?'),
        ('Tissue Specifications', '# slides cut', '# slides cut'),
        ('Tissue Specifications', 'positively charged slides required for sectioning',
         'Are positively charged slides required for sectioning?'),
        ('Tissue Specifications', 'block be returned to site after sectioning',
         'When should block be returned to site after sectioning?'),
        ('Tissue Specifications', 'slides be baked after sectioning',
         'Should slides be baked after sectioning?'),
        ('Slides', '# slides requested from sites', '# slides requested from sites'),
        ('Slides', 'slides stored or shipped to ref lab', 'Are slides stored or shipped to ref lab?'),
        ('Slides', 'slides stored and shipped, how many',
         'If slides stored and shipped, how many stored vs shipped?'),
        ('Slides', 'storage temp/condition', 'Slide storage temp/condition'),
        ('Slides', 'frequency if shipped', 'Frequency if shipped'),
        ('Slides', 'fresh frozen tissue samples be received', 'Will fresh frozen tissue samples be received?'),
        ('Slides', 'expected fixative', 'Expected fixative'),
        ('Slides', 'is local path report sent with samples', 'Is local path report sent with samples?'),
    ]

    def _populate_ap():
        got = []
        try:
            from oncology_biopsy_extractor import extract_oncology_biopsy_info
            _ap_result = extract_oncology_biopsy_info(PT + '\n\n' + DT, source_file='protocol+design')
            _mapping = dict(_ap_result.rfp_mapping)
        except Exception as _e:
            logger.warning('oncology_biopsy_extractor failed, AP section left as template defaults: %s', _e)
            _mapping = {}

        # Two evidence items the extractor surfaces but that have no dedicated template row --
        # folded into the summary rather than dropped, so a reviewer still sees them.
        _summary = _mapping.get('Anatomic Pathology / Histology - brief summary', '')
        for _extra_key in ('Bone metastasis exclusion / acceptability note', 'Cytology / FNA acceptability note'):
            _extra_val = _mapping.get(_extra_key, '')
            if _extra_val and not _extra_val.startswith('TBC'):
                _summary = (_summary + ' ' + _extra_val).strip()
        if _summary:
            _mapping['Anatomic Pathology / Histology - brief summary'] = _summary

        for table_kw, label_kw, rfp_key in _AP_FIELD_MAP:
            value = _mapping.get(rfp_key)
            if value and _ap_fill(table_kw, label_kw, value):
                got.append(rfp_key)

        # oncology_biopsy_extractor.py doesn't cover bone decalcification (a distinct concept from
        # its own "bone metastasis exclusion" evidence) -- kept from the original, simpler check.
        if re.search(r'\bbone\b[^.\n]{0,40}(?:biopsy|tissue|marrow)|decalcif', PT + DT, re.I) and \
           _ap_fill('Slides', 'decalcification of bone', 'Yes'):
            got.append('decalcification of bone')
        return got

    # Oncology Considerations table: one cell holds two content controls back-to-back --
    # "Is this an Immuno Oncology protocol?" then "Is US Oncology involved in the study?" --
    # confirmed by direct docx inspection (both share row 1's single cell, no distinguishing tag/
    # alias, so they can only be told apart by position). Previously this was a dead call
    # (`_fill_row_ccs(doc.tables[T0], 'Immuno Oncology protocol', ...)`) that searched the wrong
    # table -- that label text only exists here, in table index 2, never in T0 -- so it silently
    # matched nothing and did nothing.
    _t2 = _find_table(doc, 'Oncology Considerations')
    if _t2 is not None and len(_t2.rows) > 1:
        try:
            _t2_cell = _t2.rows[1].cells[0]
        except IndexError:
            _t2_cell = None
        if _t2_cell is not None:
            _immuno_ans = str(answers.get('immuno_oncology_override', '')).strip().lower()
            if _immuno_ans in ('yes', 'no'):
                set_nth_content_control(_t2_cell, 0, 'Yes' if _immuno_ans == 'yes' else 'No')
            rec('Is this an Immuno Oncology protocol?',
                'Yes' if _immuno_ans == 'yes' else 'No' if _immuno_ans == 'no' else None,
                'user intake (Study Details panel)', 'filled' if _immuno_ans in ('yes', 'no') else 'review')
            _us_oncology = IS_ONCOLOGY and any(
                _c.strip().lower() in ('us', 'usa', 'u.s.', 'u.s.a.') or 'united states' in _c.lower()
                for _c in COUNTRIES
            )
            set_nth_content_control(_t2_cell, 1, 'Yes' if _us_oncology else 'No')
            rec('Is US Oncology involved in the study?', 'Yes' if _us_oncology else 'No',
                'computed (oncology + US in countries)', 'computed')

    removed = []
    if ANSWERS_DICT['pediatrics'] is False and _del_rows(doc, 'Pediatric Considerations', 3):
        removed.append('Pediatric Considerations (4 rows)')
    if not ANSWERS_DICT['decentralized'] and _del_table(doc, 'Decentralized Trials'):
        removed.append('Decentralized Trials')
    if not ANSWERS_DICT['penalties_incentives']:
        for kw in ('Penalties and Incentives', 'Metrics in-Scope'):
            if _del_table(doc, kw): removed.append(kw)
    if ANSWERS_DICT['anatomic_pathology']:
        got = _populate_ap()
        removed.append(f"Anatomic Pathology kept + populated ({', '.join(got) or 'no determinable fields'})")
    else:
        for kw in ('Anatomic Pathology', 'Tissue Specifications', 'Slides'):
            if _del_table(doc, kw): removed.append(kw)

    hep = _find_table(doc, 'Calculations (Standard)')
    if hep is not None:
        def _hep_type(txt):
            t = txt.lower()
            if 'non-oncology' in t[:20]: return 'Non-oncology'
            if 'without immune checkpoint' in t: return 'Oncology — without ICI'
            if 'with immune checkpoint' in t: return 'Oncology — with ICI'
            return None
        for r in list(hep.rows)[1:]:
            ty = _hep_type(r.cells[0].text)
            if ty is not None and ty != ANSWERS_DICT['hepatic_calc']:
                r._tr.getparent().remove(r._tr)
        removed.append(f"Hepatic algorithms (kept: {ANSWERS_DICT['hepatic_calc']})")
    rec('Conditional sections (intake)', '; '.join(removed) or 'none removed', 'intake answers', 'filled')

    # Schema validation
    try:
        from schema import validate_extraction
        extraction_data = {
            'general_information': {
                'lilly_requestor_name': 'Lisa Brennan',
                'lilly_requestor_email': 'lisa.brennan@lilly.com',
            },
            'protocol_information': {
                'protocol_alias': alias or '',
                'protocol_title': title or '',
                'compound': compound_val or '',
                'phase': phase_val or '',
                'therapeutic_area': ta or '',
            },
        }
        schema_errors = validate_extraction(extraction_data)
        if schema_errors:
            for err in schema_errors:
                rec('Schema validation', err, 'schema', 'review')
                logger.warning('Schema validation: %s', err)
    except Exception as _schema_err:
        logger.warning('Schema validation skipped: %s', _schema_err)

    doc.save(str(output_path))

    # ---- Report ---
    filled  = sum(1 for f in findings if f.status == 'filled')
    comp    = sum(1 for f in findings if f.status == 'computed')
    review  = sum(1 for f in findings if f.status == 'review')
    report_lines = []
    report_lines.append('# Central Laboratory RFP — Fill Report\n\n')
    report_lines.append(f'Generated {submitted}. Sources: protocol {alias or "(unresolved)"} '
                        f'+ clinical design elements{" + " + Path(previous_rfp_path).name if previous_rfp_path else ""}.\n\n')
    report_lines.append(f'**Coverage:** {filled} filled · {comp} computed · {review} need review '
                        f'(of {len(findings)} tracked items)\n\n')
    report_lines.append('| Field | Value | Source | Status |\n|---|---|---|---|\n')
    for entry in findings:
        # Every cell needs escaping, not just value -- a couple of field names (e.g. "Analytes —
        # Appendix 2 table (Test | Comments)") contain a literal pipe themselves, which breaks
        # any GFM pipe-table parser reading this file back (confirmed: it shifts that row's
        # remaining columns by one).
        ff = str(entry.field).replace('|', '/')
        vv = '' if entry.value is None else str(entry.value).replace('|', '/')[:80]
        ss = str(entry.source).replace('|', '/')
        report_lines.append(f'| {ff} | {vv} | {ss} | {entry.status} |\n')
    report_lines.append('\n## Notes\n')
    report_lines.append('- **Analytes:** Appendix table rendered from the protocol markdown pipe table '
                        '(Section 10 — Clinical Laboratory Tests | Comments).\n')
    report_lines.append('- **Schedule of Activities:** Table 7 rebuilt from the protocol markdown pipe table '
                        '(Section 1 — Clinical Laboratory Tests).\n')
    report_lines.append('- **Specimen Management:** Referral (LTS PK / LTS Immunogenicity) and Storage '
                        f'(LTS DNA/Serum/Plasma/RNA) filled from {_specimen_source} into editable table cells.\n')
    report_lines.append('- Items marked **review** are absent from all sources — left as visible tokens, never guessed.\n')
    report_text = ''.join(report_lines)

    Path(report_path).write_text(report_text, encoding='utf-8')

    print(f'Saved: {Path(output_path).name}')
    print(f'Coverage: {filled} filled, {comp} computed, {review} review (of {len(findings)})')
    print(f'Analyte table rows: {len(ana_rows)} (2-column Test | Comments)')

    return Report(findings=findings, filled=filled, computed=comp, review_count=review,
                  report_text=report_text, output_path=output_path,
                  countries=CTRIES, soa_headers=soa_visits, soa_rows=soa_rows)


# ======================================================================
# CLI entry point (backward-compatible with env-var configuration)
# ======================================================================

if __name__ == '__main__':
    INPUT   = Path(os.environ.get('RFP_INPUT_DIR', '/mnt/workspace/input'))
    OUTPUT  = Path(os.environ.get('RFP_OUTPUT_DIR', '/mnt/workspace/output'))
    OUTPUT.mkdir(parents=True, exist_ok=True)
    PROTOCOL = Path(os.environ.get('RFP_PROTOCOL', str(INPUT / 'OIAH protocol markdown.txt')))
    DESIGN   = Path(os.environ.get('RFP_DESIGN',   str(INPUT / 'OIAH Clinical Design Elements markdown.txt')))
    TEMPLATE = Path(os.environ.get('RFP_TEMPLATE', str(INPUT / 'Central_Laboratory_RFP_PlainText.docx')))
    OUTDOC   = Path(os.environ.get('RFP_OUTDOC',   str(OUTPUT / 'Central_Laboratory_RFP_OIAH_populated.docx')))
    REPORT   = Path(os.environ.get('RFP_REPORT',   str(OUTPUT / 'RFP_fill_report.md')))
    ANSWERS  = json.loads(os.environ.get('RFP_ANSWERS', '{}'))
    PROTOCOL_PDF = os.environ.get('RFP_PROTOCOL_PDF', '')
    PREVIOUS_RFP = os.environ.get('RFP_PREVIOUS', '')

    pt_text = PROTOCOL.read_text(errors='replace', encoding='utf-8')
    dt_text = DESIGN.read_text(errors='replace', encoding='utf-8')

    main(protocol_text=pt_text, design_text=dt_text,
         template_path=str(TEMPLATE), output_path=str(OUTDOC),
         report_path=str(REPORT), answers=ANSWERS,
         protocol_pdf_path=PROTOCOL_PDF, previous_rfp_path=PREVIOUS_RFP)
