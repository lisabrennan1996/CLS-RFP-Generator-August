"""
oncology_biopsy_extractor.py

Purpose:
    Extract oncology biopsy / tumor tissue requirements from protocol text,
    PDF, or DOCX files and map them into Central Lab RFP Anatomic Pathology /
    Histology fields.

Typical protocol source language this targets:
    - Biomarker Assessments in Tumor Samples
    - Tumor tissue samples
    - Archival / archived tumor tissue
    - FFPE / formalin-fixed paraffin-embedded
    - Tissue blocks
    - Unstained slides
    - Tumor content
    - Bone metastasis unsuitable / not acceptable
    - Pathology / molecular reports
    - Optional tissue at disease progression

Outputs:
    - A structured dictionary of extracted evidence
    - A suggested RFP mapping dictionary
    - Optional JSON file output

Dependencies:
    pip install python-docx pdfplumber

Notes:
    - The script is intentionally conservative.
    - If something is not explicitly found, it returns "TBC / per Lab Manual"
      rather than guessing.
"""

from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, List, Optional, Tuple


# -----------------------------
# Optional file readers
# -----------------------------

def read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX parsing. Install with: pip install python-docx"
        ) from exc

    doc = Document(str(path))
    parts: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Pull text from tables too, because protocols and appendices often use tables.
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise ImportError(
            "pdfplumber is required for PDF parsing. Install with: pip install pdfplumber"
        ) from exc

    parts: List[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"\n\n--- PAGE {i} ---\n{text}")
    return "\n".join(parts)


def read_any(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return read_txt(path)
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".pdf":
        return read_pdf(path)

    raise ValueError(f"Unsupported file type: {suffix}. Use .txt, .docx, or .pdf")


# -----------------------------
# Data structures
# -----------------------------

@dataclass
class EvidenceItem:
    label: str
    value: str
    evidence: str
    confidence: str


@dataclass
class OncologyBiopsyExtraction:
    source_file: str
    relevant_section_heading: str
    section_excerpt: str
    evidence_items: List[EvidenceItem]
    rfp_mapping: Dict[str, str]
    flags_for_manual_review: List[str]


# -----------------------------
# Text helpers
# -----------------------------

def normalize_text(text: str) -> str:
    # Keep original-ish line breaks but normalize whitespace enough for regex.
    text = text.replace(" ", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def clean_excerpt(text: str, max_len: int = 900) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) <= max_len:
        return text
    return text[:max_len].rstrip() + "..."


def trim_to_sentence(text: str, max_len: int = 150) -> str:
    """Cut a matched span down to its first complete sentence (or max_len characters,
    whichever is shorter). Some patterns have wide trailing-context windows (`.{0,250}`)
    with nothing anchoring what follows, so the match always consumes the full window --
    including, in practice, unrelated text past the sentence the keyword actually appears
    in. Used for "value" fields that should read as one short answer, not a running quote;
    the raw match is still kept separately as the evidence snippet."""
    text = re.sub(r"\s+", " ", text).strip()
    period_idx = text.find(". ")
    if period_idx != -1 and period_idx + 1 <= max_len:
        return text[:period_idx + 1].strip()
    return clean_excerpt(text, max_len)


def find_first(patterns: List[str], text: str, flags: int = re.IGNORECASE) -> Optional[re.Match]:
    for pattern in patterns:
        match = re.search(pattern, text, flags)
        if match:
            return match
    return None


def sentence_windows(text: str, keywords: List[str], window_chars: int = 500) -> List[str]:
    """
    Return text windows around relevant keywords.
    Useful for evidence snippets.
    """
    hits: List[str] = []
    lowered = text.lower()

    for kw in keywords:
        start = 0
        kw_l = kw.lower()
        while True:
            idx = lowered.find(kw_l, start)
            if idx == -1:
                break
            left = max(0, idx - window_chars)
            right = min(len(text), idx + len(kw) + window_chars)
            snippet = clean_excerpt(text[left:right], max_len=1100)
            hits.append(snippet)
            start = idx + len(kw)

    # Deduplicate similar snippets.
    deduped: List[str] = []
    seen = set()
    for h in hits:
        key = h[:150].lower()
        if key not in seen:
            seen.add(key)
            deduped.append(h)

    return deduped[:8]


# -----------------------------
# Section detection
# -----------------------------

SECTION_HEADINGS = [
    r"biomarker assessments in tumor samples",
    r"exploratory biomarker assessments in tumor samples",
    r"tumou?r tissue samples?",
    r"archiv(?:al|ed) tumou?r tissue",
    r"screening tumou?r samples?",
    r"sample at time of disease progression",
    r"tumou?r biopsy",
    r"biopsy of metastatic lesion",
]


def extract_relevant_section(text: str) -> Tuple[str, str]:
    """
    Try to isolate the highest-value oncology tumor tissue section.
    If no clean section is found, return windows around keywords.
    """
    text_norm = normalize_text(text)

    # Look for a heading, then take until the next numbered section-like heading.
    heading_regex = re.compile(
        r"(?P<heading>(?:\d+(?:\.\d+)*\.?\s*)?"
        r"(?:Exploratory\s+)?Biomarker Assessments in Tumou?r Samples|"
        r"(?:\d+(?:\.\d+)*\.?\s*)?Screening Tumou?r Samples|"
        r"(?:\d+(?:\.\d+)*\.?\s*)?Tumou?r Tissue Samples|"
        r"(?:\d+(?:\.\d+)*\.?\s*)?Sample at Time of Disease Progression)"
        r"(?P<body>.*?)(?=\n\s*\d+(?:\.\d+){1,4}\.?\s+[A-Z]|\n\s*[A-Z][A-Za-z ]{5,80}\n|$)",
        re.IGNORECASE | re.DOTALL,
    )

    matches = list(heading_regex.finditer(text_norm))
    if matches:
        # Pick the longest likely section.
        best = max(matches, key=lambda m: len(m.group("body")))
        heading = clean_excerpt(best.group("heading"), 160)
        body = clean_excerpt(best.group("heading") + " " + best.group("body"), 2500)
        return heading, body

    # Fallback: keyword windows.
    keywords = [
        "Biomarker Assessments in Tumor Samples",
        "Biomarker Assessments in Tumour Samples",
        "tumor tissue",
        "tumour tissue",
        "archival tissue",
        "archived tumor tissue",
        "FFPE",
        "unstained slides",
        "disease progression",
    ]
    windows = sentence_windows(text_norm, keywords, window_chars=700)
    if windows:
        return "Keyword-based tumor tissue excerpts", "\n\n---\n\n".join(windows)

    return "No clear tumor tissue section found", ""


# -----------------------------
# Extractors
# -----------------------------

def evidence(label: str, value: str, snippet: str, confidence: str = "medium") -> EvidenceItem:
    return EvidenceItem(
        label=label,
        value=value,
        evidence=clean_excerpt(snippet, 700),
        confidence=confidence,
    )


def extract_oncology_biopsy_info(text: str, source_file: str = "") -> OncologyBiopsyExtraction:
    text_norm = normalize_text(text)
    heading, section = extract_relevant_section(text_norm)
    search_space = section if section else text_norm

    items: List[EvidenceItem] = []
    flags: List[str] = []

    # Oncology/tumor tissue presence
    m = find_first([
        r"\b(?:tumou?r tissue samples?|archiv(?:al|ed) tumou?r tissue|biomarker assessments in tumou?r samples|tumou?r biopsy)\b"
    ], search_space)
    if m:
        items.append(evidence(
            "Tumor tissue / biopsy language present",
            "Yes",
            search_space[max(0, m.start() - 250): m.end() + 450],
            "high",
        ))

    # Source/sample collection methods
    m = find_first([
        r"(?:may include|include[s]?|collection methods?.{0,80})(?:surgical resection|core(?: needle)? biops|excisional|incisional|endobronchial|fine needle aspiration|cell blocks?).{0,350}",
        r"(?:surgical resection|core(?: needle)? biops|excisional|incisional|endobronchial|fine needle aspiration|cell blocks?).{0,350}",
    ], search_space)
    if m:
        items.append(evidence(
            "Acceptable collection methods / sample types",
            clean_excerpt(m.group(0), 350),
            m.group(0),
            "high",
        ))

    # FFPE
    m = find_first([
        r"(?:FFPE|formalin-fixed paraffin-embedded|formalin fixed paraffin embedded).{0,250}",
        r".{0,120}(?:FFPE|formalin-fixed paraffin-embedded|formalin fixed paraffin embedded).{0,150}",
    ], search_space)
    if m:
        items.append(evidence(
            "FFPE / fixation requirement",
            clean_excerpt(m.group(0), 260),
            m.group(0),
            "high",
        ))

    # Blocks preferred
    m = find_first([
        r"(?:blocks? (?:are )?preferred|recommended that .*?blocks? be submitted|FFPE .*?blocks? be submitted).{0,250}",
        r".{0,150}(?:tissue blocks? are preferred|tumou?r blocks? is unavailable|tumor blocks? is unavailable).{0,250}",
    ], search_space)
    if m:
        items.append(evidence(
            "Tissue block preference",
            clean_excerpt(m.group(0), 300),
            m.group(0),
            "high",
        ))

    # Unstained slides count and thickness -- value is built from the regex's own named capture
    # groups (count / thickness) rather than the whole matched span, which otherwise dragged in
    # up to 180-300 characters of unrelated trailing protocol text as the "value" (fine for the
    # separately-kept evidence snippet below, not fine for a field meant to hold just a number).
    slide_patterns = [
        r"(?:approximately|minimum of|at least)?\s*(?P<count>\d{1,3})\s+(?:serially cut\s+)?unstained slides?.{0,180}",
        r"unstained slides?.{0,80}(?P<thickness>\d+(?:\s*to\s*\d+)?\s*(?:microns?|µm|um)).{0,160}",
        r"slides?.{0,80}(?:cut|sectioned)\s*(?P<thickness2>\d+(?:\s*to\s*\d+)?\s*(?:microns?|µm|um)).{0,160}",
    ]
    m = find_first(slide_patterns, search_space)
    if m:
        gd = m.groupdict()
        _thickness = gd.get('thickness') or gd.get('thickness2')
        _parts = []
        if gd.get('count'):
            _parts.append(f"{gd['count']} unstained slides")
        if _thickness:
            _parts.append(f"{_thickness} thick")
        _slide_value = ', '.join(_parts) if _parts else clean_excerpt(m.group(0), 300)
        items.append(evidence(
            "Unstained slide requirements",
            _slide_value,
            m.group(0),
            "high",
        ))

    # Tumor content
    m = find_first([
        r"(?:preferred\s+)?tumou?r content.{0,80}(?:at least|>=|≥|minimum of)?\s*\d{1,3}\s*%",
        r"(?:verification of|verify).{0,80}\d{1,3}\s*%\s*tumou?r content",
        r"\d{1,3}\s*%\s*tumou?r content.{0,120}",
    ], search_space)
    if m:
        items.append(evidence(
            "Tumor content requirement",
            clean_excerpt(m.group(0), 220),
            m.group(0),
            "high",
        ))

    # Bone metastasis exclusion
    m = find_first([
        r"(?:bone metastasis|bone metastases).{0,180}(?:unsuitable|not acceptable|are not acceptable|not acceptable)",
        r"(?:unsuitable|not acceptable).{0,180}(?:bone metastasis|bone metastases)",
    ], search_space)
    if m:
        items.append(evidence(
            "Bone metastasis exclusion",
            clean_excerpt(m.group(0), 260),
            m.group(0),
            "high",
        ))

    # Cytology / FNA acceptability
    m = find_first([
        r"(?:cell blocks?|cytology|fine needle aspiration|FNA).{0,220}(?:acceptable|not acceptable|unsuitable)",
        r"(?:acceptable|not acceptable|unsuitable).{0,220}(?:cell blocks?|cytology|fine needle aspiration|FNA)",
    ], search_space)
    if m:
        items.append(evidence(
            "Cytology / FNA handling",
            clean_excerpt(m.group(0), 300),
            m.group(0),
            "medium",
        ))

    # Pathology / molecular reports -- value is trimmed to the sentence the match actually
    # appears in rather than the raw match (the trailing `.{0,250}` on each pattern has nothing
    # anchoring it, so it always consumed the full 250 characters, which in practice ran past the
    # end of this section into unrelated text).
    m = find_first([
        r"(?:deidentified|de-identified|coded).{0,250}(?:pathology|molecular) reports?.{0,250}",
        r"(?:pathology|molecular) reports?.{0,250}(?:requested|submitted|accompany).{0,250}",
        r"personal identifiers.{0,220}(?:removed|must be removed).{0,220}",
    ], search_space)
    if m:
        items.append(evidence(
            "Pathology / molecular report requirement",
            trim_to_sentence(m.group(0), 150),
            m.group(0),
            "high",
        ))

    # Screening availability confirmation
    m = find_first([
        r"confirmation of.{0,120}(?:archiv(?:al|ed) )?tumou?r tissue.{0,160}(?:screening|required|required for all participants)",
        r"availability of.{0,120}tumou?r tissue.{0,160}(?:screening|required|required for all participants)",
    ], search_space)
    if m:
        items.append(evidence(
            "Tumor tissue availability confirmation",
            clean_excerpt(m.group(0), 300),
            m.group(0),
            "high",
        ))

    # Optional progression biopsy
    m = find_first([
        r"optional.{0,80}tumou?r tissue.{0,220}(?:progression|disease progression)",
        r"(?:progression|disease progression).{0,220}optional.{0,80}tumou?r tissue",
        r"biopsy of metastatic lesion.{0,260}disease progression",
    ], search_space)
    if m:
        items.append(evidence(
            "Optional progression tissue / biopsy",
            clean_excerpt(m.group(0), 350),
            m.group(0),
            "high",
        ))

    # Storage / retention
    m = find_first([
        r"(?:samples?|tissue samples?).{0,120}(?:retained|stored|storage).{0,220}",
        r"(?:retained|stored|storage).{0,220}(?:samples?|tissue samples?)",
    ], search_space)
    if m:
        items.append(evidence(
            "Storage / retention language",
            clean_excerpt(m.group(0), 300),
            m.group(0),
            "medium",
        ))

    # Build RFP mapping conservatively.
    rfp_mapping = build_rfp_mapping(items)

    # Manual review flags.
    if not any(i.label == "Unstained slide requirements" for i in items):
        flags.append("Slide count/thickness not explicitly found; leave slide details as TBC / per Lab Manual.")
    if not any(i.label == "Storage / retention language" for i in items):
        flags.append("Storage/return instructions not explicitly found; leave storage/return as TBC / per Lab Manual.")
    if not any(i.label == "Pathology / molecular report requirement" for i in items):
        flags.append("Pathology report requirement not explicitly found; confirm whether deidentified local pathology report is required.")
    if not any(i.label == "Tumor content requirement" for i in items):
        flags.append("Tumor content threshold not explicitly found; confirm whether % tumor assessment/H&E review is required.")
    if not items:
        flags.append("No oncology biopsy/tumor tissue language found using current search patterns.")

    return OncologyBiopsyExtraction(
        source_file=source_file,
        relevant_section_heading=heading,
        section_excerpt=section,
        evidence_items=items,
        rfp_mapping=rfp_mapping,
        flags_for_manual_review=flags,
    )


def get_item(items: List[EvidenceItem], label: str) -> Optional[EvidenceItem]:
    for item in items:
        if item.label == label:
            return item
    return None


def build_rfp_mapping(items: List[EvidenceItem]) -> Dict[str, str]:
    """
    Translate extracted evidence into the RFP Anatomic Pathology / Histology section.
    Do not invent anything. If not found, return TBC / per Lab Manual.
    """
    tumor_present = get_item(items, "Tumor tissue / biopsy language present")
    sample_types = get_item(items, "Acceptable collection methods / sample types")
    ffpe = get_item(items, "FFPE / fixation requirement")
    blocks = get_item(items, "Tissue block preference")
    slides = get_item(items, "Unstained slide requirements")
    tumor_content = get_item(items, "Tumor content requirement")
    bone = get_item(items, "Bone metastasis exclusion")
    cytology = get_item(items, "Cytology / FNA handling")
    reports = get_item(items, "Pathology / molecular report requirement")
    availability = get_item(items, "Tumor tissue availability confirmation")
    progression = get_item(items, "Optional progression tissue / biopsy")
    storage = get_item(items, "Storage / retention language")

    summary_parts = []
    if tumor_present:
        summary_parts.append("Tumor tissue samples are described for biomarker assessment.")
    if blocks:
        summary_parts.append(blocks.value)
    if slides:
        summary_parts.append(slides.value)
    if progression:
        summary_parts.append("Optional tumor tissue collection at disease progression is described.")
    if bone:
        summary_parts.append(bone.value)

    summary = " ".join(summary_parts) if summary_parts else "TBC - no explicit oncology tumor tissue language found."

    archived_or_fresh = "TBC / per Protocol and Lab Manual"
    if availability:
        archived_or_fresh = "Archived tumor tissue availability is referenced during screening."
    if progression:
        archived_or_fresh += " Optional new tissue at progression may also apply."

    block_slide_value = "TBC / per Lab Manual"
    if blocks and slides:
        block_slide_value = f"{blocks.value} If block unavailable: {slides.value}"
    elif blocks:
        block_slide_value = blocks.value
    elif slides:
        block_slide_value = slides.value

    return {
        # Main AP/Histo summary field
        "Anatomic Pathology / Histology - brief summary":
            summary,

        # Tissue Specifications
        "Will submitted samples be archived or fresh? If combo, specify percentage of each.":
            archived_or_fresh,

        "Is % tumor assessment required on an H&E slide for block/slides submitted?":
            tumor_content.value if tumor_content else "TBC - protocol does not explicitly state H&E/% tumor assessment requirement in extracted text.",

        "Is sample inspection required on blocks/slides received from site before storage?":
            "Yes - recommended to verify receipt of expected blocks/slides and adequacy, but confirm against Lab Manual / vendor specs.",

        "H&E level of complexity":
            "TBC - oncology tumor tissue review likely requires pathology/biomarker adequacy review if specified; confirm complexity with vendor/pathology team.",

        # Sectioning of Blocks
        "Block / slide submission requirement":
            block_slide_value,

        "If multiple blocks are received, is H&E required on each block?":
            "TBC / per Lab Manual",

        "Should slides be cut upon receipt or upon request?":
            "TBC / per Lab Manual",

        "Are tissue curls required?":
            "TBC / per Lab Manual",

        "# slides cut":
            slides.value if slides else "TBC / per Lab Manual",

        "Are positively charged slides required for sectioning?":
            "TBC / per Lab Manual",

        "When should block be returned to site after sectioning?":
            "TBC / per Lab Manual",

        "Should slides be baked after sectioning?":
            "TBC / per Lab Manual",

        # Slides
        "# slides requested from sites":
            slides.value if slides else "TBC / per Lab Manual",

        "Are slides stored or shipped to ref lab?":
            "TBC - depends whether central lab performs testing or forwards to referral/specialty lab.",

        "If slides stored and shipped, how many stored vs shipped?":
            "TBC / per Lab Manual",

        "Slide storage temp/condition":
            "TBC / per Lab Manual",

        "Frequency if shipped":
            "TBC / per Lab Manual",

        # Fresh Tissue
        "Will fresh frozen tissue samples be received?":
            "TBC - extracted oncology examples primarily reference FFPE/archival tumor tissue unless protocol states fresh tissue.",

        "Expected fixative":
            ffpe.value if ffpe else "TBC / per Lab Manual",

        "Is local path report sent with samples?":
            reports.value if reports else "TBC - confirm whether deidentified pathology/molecular report is required.",

        "Bone metastasis exclusion / acceptability note":
            bone.value if bone else "TBC - not explicitly found.",

        "Cytology / FNA acceptability note":
            cytology.value if cytology else "TBC - not explicitly found.",

        "Storage / retention":
            storage.value if storage else "TBC / per Lab Manual",
    }


# -----------------------------
# Markdown reporting
# -----------------------------

def to_markdown(result: OncologyBiopsyExtraction) -> str:
    lines: List[str] = []

    lines.append("# Oncology Biopsy / Tumor Tissue Extraction")
    lines.append("")
    lines.append(f"**Source file:** {result.source_file or 'N/A'}")
    lines.append(f"**Relevant section detected:** {result.relevant_section_heading}")
    lines.append("")

    if result.section_excerpt:
        lines.append("## Relevant excerpt")
        lines.append("")
        lines.append(result.section_excerpt)
        lines.append("")

    lines.append("## Evidence found")
    lines.append("")
    if result.evidence_items:
        for item in result.evidence_items:
            lines.append(f"### {item.label}")
            lines.append(f"- **Value:** {item.value}")
            lines.append(f"- **Confidence:** {item.confidence}")
            lines.append(f"- **Evidence:** {item.evidence}")
            lines.append("")
    else:
        lines.append("No direct biopsy / tumor tissue evidence found.")
        lines.append("")

    lines.append("## Suggested RFP mapping")
    lines.append("")
    for key, value in result.rfp_mapping.items():
        lines.append(f"### {key}")
        lines.append(value)
        lines.append("")

    if result.flags_for_manual_review:
        lines.append("## Manual review flags")
        lines.append("")
        for flag in result.flags_for_manual_review:
            lines.append(f"- {flag}")
        lines.append("")

    return "\n".join(lines)


# -----------------------------
# CLI
# -----------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Extract oncology biopsy / tumor tissue requirements from protocol files."
    )
    parser.add_argument(
        "input_file",
        help="Path to protocol file: .pdf, .docx, or .txt"
    )
    parser.add_argument(
        "--json-out",
        default=None,
        help="Optional path to write structured JSON output."
    )
    parser.add_argument(
        "--md-out",
        default=None,
        help="Optional path to write Markdown report."
    )
    parser.add_argument(
        "--print-md",
        action="store_true",
        help="Print Markdown report to stdout."
    )

    args = parser.parse_args()

    input_path = Path(args.input_file)
    text = read_any(input_path)

    result = extract_oncology_biopsy_info(
        text=text,
        source_file=input_path.name
    )

    result_dict = asdict(result)

    if args.json_out:
        Path(args.json_out).write_text(
            json.dumps(result_dict, indent=2, ensure_ascii=False),
            encoding="utf-8"
        )

    md = to_markdown(result)

    if args.md_out:
        Path(args.md_out).write_text(md, encoding="utf-8")

    if args.print_md or not args.json_out and not args.md_out:
        print(md)


if __name__ == "__main__":
    main()
